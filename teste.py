import os
import re
import shutil
import subprocess
from datetime import datetime, date
from urllib.parse import unquote
from flask import Flask, request, jsonify, render_template, send_file
from docx import Document
import tempfile
import json

app = Flask(__name__)

# --- CONFIGURAÇÕES ---
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
UPLOAD_FOLDER = os.path.join(BASE_DIR, 'uploads')
PROCESSED_FOLDER = os.path.join(UPLOAD_FOLDER, 'concluidos')
INTERNACAO_FOLDER = os.path.join(UPLOAD_FOLDER, 'internacao')
PRE_FICHAS_FOLDER = os.path.join(BASE_DIR, 'pre_fichas')
MODELS_FOLDER = os.path.join(BASE_DIR, 'modelos')
BIN_FOLDER = os.path.join(BASE_DIR, 'bin')

# --- CONTROLE DE VERSÃO ---
APP_HEMO_VERSION = "9.0.3"
APP_HEMO_FILE = "AppHemo.jar"
APP_RECEPCAO_VERSION = "5.1"
APP_RECEPCAO_FILE = "AppRecepcao.jar"

EXTENSOES_PERMITIDAS = ('.doc', '.docx')

for pasta in [UPLOAD_FOLDER, PROCESSED_FOLDER, INTERNACAO_FOLDER, PRE_FICHAS_FOLDER, MODELS_FOLDER, BIN_FOLDER]:
    os.makedirs(pasta, exist_ok=True)

# --- FUNÇÕES AUXILIARES ---

def extrair_dados_completos(caminho_docx):
    dados = {
        "nome": "Nome não identificado",
        "procedencia": "Não informada",
        "tipo_procedimento": "Não identificado"
    }
    try:
        doc = Document(caminho_docx)
        texto_completo = "\n".join(p.text for p in doc.paragraphs)
        for t in doc.tables:
            for r in t.rows:
                for c in r.cells:
                    texto_completo += "\n" + c.text

        # (Ponto 1) Regex para nome corrigido com lógica mais segura
        match_linha_nome = re.search(r'^(.*(NOME|PACIENTE)\s*:.*)$', texto_completo, re.MULTILINE | re.IGNORECASE)
        if match_linha_nome:
            linha_completa = match_linha_nome.group(1)
            # Pega o que vem depois do último ':' na linha
            nome_bruto = linha_completa.split(':')[-1]
            dados["nome"] = " ".join(nome_bruto.strip().split())

        texto_upper = texto_completo.upper()
        match_proc = re.search(r'(?:PROCED.NCIA|ORIGEM|UNIDADE DE ORIGEM)\s*[:\s_]+(.*?)(?:\n|M.DICO|CONV.NIO|LEITO|DATA|HORA|SETOR|PACIENTE)', texto_completo, re.IGNORECASE | re.DOTALL)
        if match_proc:
            dados["procedencia"] = match_proc.group(1).strip().replace('\n', ' ')

        if "ANGIOPLASTIA" in texto_upper or "PTCA" in texto_upper:
            dados["tipo_procedimento"] = "ANGIOPLASTIA"
        elif "CATETERISMO" in texto_upper:
            dados["tipo_procedimento"] = "CATETERISMO"

        return dados
    except Exception as e:
        print(f"Erro ao extrair dados de {os.path.basename(caminho_docx)}: {e}")
        return dados

def gerar_nome_seguro(nome_paciente, pasta_destino):
    data_hoje = datetime.now().strftime('%Y%m%d')
    nome_base = f"{data_hoje}_{nome_paciente.replace(' ', '_')}"
    nome_base = re.sub(r'[\\/*?:"<>|]', "", nome_base)
    caminho_final = os.path.join(pasta_destino, f"{nome_base}.docx")
    contador = 1
    while os.path.exists(caminho_final):
        caminho_final = os.path.join(pasta_destino, f"{nome_base}_{contador}.docx")
        contador += 1
    return os.path.basename(caminho_final)

def converter_doc_para_docx(caminho_doc):
    pasta = os.path.dirname(caminho_doc)
    try:
        subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "docx", caminho_doc, "--outdir", pasta],
            stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, timeout=30, check=True
        )
        novo_caminho = caminho_doc.rsplit('.', 1)[0] + '.docx'
        return novo_caminho if os.path.exists(novo_caminho) else None
    except (subprocess.CalledProcessError, FileNotFoundError, subprocess.TimeoutExpired):
        return None

def substituir_placeholders_py(doc, data):
    for p in doc.paragraphs:
        for r in p.runs:
            for key, value in data.items():
                if key in r.text:
                    r.text = r.text.replace(key, str(value))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        for key, value in data.items():
                            if key in r.text:
                                r.text = r.text.replace(key, str(value))

# --- ROTAS PRINCIPAIS E DE API ---

@app.route('/', methods=['GET', 'POST'])
def index():
    enviados, erros = [], []
    if request.method == 'POST':
        arquivos = request.files.getlist('files')
        for f in arquivos:
            if not f.filename.lower().endswith(EXTENSOES_PERMITIDAS):
                erros.append(f"{f.filename}: Extensão não permitida")
                continue
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                f.save(tmp.name)
                caminho_tmp = tmp.name
            caminho_docx = caminho_tmp
            if f.filename.lower().endswith('.doc'):
                caminho_docx = converter_doc_para_docx(caminho_tmp)
                if not caminho_docx:
                    erros.append(f"{f.filename}: Falha ao converter .doc para .docx")
                    os.remove(caminho_tmp)
                    continue
            dados_paciente = extrair_dados_completos(caminho_docx)
            nome_final = gerar_nome_seguro(dados_paciente["nome"], UPLOAD_FOLDER)
            caminho_destino = os.path.join(UPLOAD_FOLDER, nome_final)
            shutil.move(caminho_docx, caminho_destino)
            enviados.append(f"{f.filename} -> {nome_final}")
            if f.filename.lower().endswith('.doc'):
                os.remove(caminho_tmp)
    return render_template('index.html', enviados=enviados, erros=erros)

@app.route('/historico')
def historico():
    pendentes, concluidos, em_processamento = [], []
    for f in sorted(os.listdir(PROCESSED_FOLDER), reverse=True):
        if f.endswith('.docx'):
            concluidos.append(f)
    for f in sorted(os.listdir(UPLOAD_FOLDER), reverse=True):
        if f.endswith('.docx'):
            caminho_completo = os.path.join(UPLOAD_FOLDER, f)
            stat = os.stat(caminho_completo)
            if stat.st_mtime > stat.st_ctime + 1:
                em_processamento.append(f)
            else:
                pendentes.append(f)
    return render_template('historico.html', pendentes=pendentes, concluidos=concluidos, em_processamento=em_processamento)

@app.route('/view/<path:filename>')
def view_file(filename):
    filename = unquote(filename)
    for pasta in [UPLOAD_FOLDER, PROCESSED_FOLDER, INTERNACAO_FOLDER]:
        caminho = os.path.join(pasta, filename)
        if os.path.exists(caminho):
            try:
                doc = Document(caminho)
                texto_completo = "\n".join(p.text for p in doc.paragraphs)
                return render_template('view.html', filename=filename, content=texto_completo)
            except Exception as e:
                return f"Erro ao ler o arquivo: {e}", 500
    return "Arquivo não encontrado", 404

@app.route('/api/pendentes')
def api_pendentes():
    lista = []
    for f in os.listdir(UPLOAD_FOLDER):
        if f.endswith('.docx'):
            dados = extrair_dados_completos(os.path.join(UPLOAD_FOLDER, f))
            lista.append({
                'arquivo': f,
                'nome': dados['nome'],
                'procedencia': dados['procedencia'],
                'tipo_procedimento': dados['tipo_procedimento']
            })
    return jsonify(lista)

@app.route('/api/internacao/listar')
def listar_internacao():
    lista = []
    hoje_str = datetime.now().strftime('%Y%m%d')
    pastas_a_verificar = [UPLOAD_FOLDER, PROCESSED_FOLDER]
    arquivos_vistos = set()

    for pasta in pastas_a_verificar:
        if not os.path.exists(pasta):
            continue
        for f in os.listdir(pasta):
            if f.startswith(hoje_str) and f.endswith('.docx') and f not in arquivos_vistos:
                caminho_completo = os.path.join(pasta, f)
                try:
                    dados = extrair_dados_completos(caminho_completo)
                    lista.append({
                        'nome': dados['nome'],
                        'procedencia': dados['procedencia'], # (Ponto 2) Adicionado
                        'arquivo': f,
                        'data': datetime.now().strftime('%d/%m/%Y'),
                        'tipo': dados['tipo_procedimento']
                    })
                    arquivos_vistos.add(f)
                except Exception as e:
                    print(f"Erro ao processar '{f}' para lista de internação: {e}")

    return jsonify(sorted(lista, key=lambda x: x['nome']))


@app.route('/api/internacao/gerar', methods=['POST'])
def gerar_docs_internacao():
    dados_req = request.get_json()
    if not dados_req:
        return jsonify({"erro": "Requisição sem JSON"}), 400

    modelos_solicitados = dados_req.get("modelos", [])
    if not modelos_solicitados:
        return jsonify({"erro": "Nenhum modelo solicitado"}), 400

    dados_placeholders = {
        "{{NOME}}": dados_req.get("nome", "").upper(),
        "{{PROCEDENCIA}}": dados_req.get("procedencia", ""), # (Ponto 2) Adicionado
        "{{DATA_HOJE}}": datetime.now().strftime('%d/%m/%Y'),
        "{{ARTERIAS}}": dados_req.get("arterias", ""),
        "{{STENTS}}": dados_req.get("stents", ""),
        "{{ARTERIAS_JUST}}": dados_req.get("arterias_just", ""),
        "{{STENTS_JUST}}": dados_req.get("stents_just", ""),
        "{{CHK_CLINICO}}": dados_req.get("chk_clinico", " "),
        "{{CHK_ANGIO}}": dados_req.get("chk_angio", " "),
        "{{CHK_CIRURGIA}}": dados_req.get("chk_cirurgia", " "),
        "{{OBS_TXT}}": dados_req.get("obs_txt", "")
    }

    arquivos_gerados = []
    for nome_modelo in modelos_solicitados:
        caminho_modelo = os.path.join(MODELS_FOLDER, nome_modelo)
        if not os.path.exists(caminho_modelo):
            continue

        try:
            doc = Document(caminho_modelo)
            substituir_placeholders_py(doc, dados_placeholders)

            nome_base, _ = os.path.splitext(nome_modelo)
            nome_paciente_safe = re.sub(r'[\\/*?:"<>|]', "", dados_req.get("nome", "paciente"))
            nome_final = f"{nome_paciente_safe}_{nome_base}.docx"

            caminho_final = os.path.join(INTERNACAO_FOLDER, nome_final)
            doc.save(caminho_final)
            arquivos_gerados.append(nome_final)
        except Exception as e:
            print(f"Erro ao gerar {nome_modelo}: {e}")

    return jsonify({"arquivos_gerados": arquivos_gerados}), 201

@app.route('/api/modelos/<path:filename>')
def baixar_modelo_generico(filename):
    caminho = os.path.join(MODELS_FOLDER, filename)
    if os.path.exists(caminho):
        return send_file(caminho, as_attachment=True)
    return "Modelo não encontrado", 404

@app.route('/api/modelo/<tipo>')
def modelo(tipo):
    nome_map = {
        'cateterismo': 'Laudo de cateterismo.docx',
        'angioplastia': 'Laudo de Angioplastia.docx'
    }
    nome_arquivo = nome_map.get(tipo)
    if not nome_arquivo:
        return "Tipo de modelo inválido", 400
    return baixar_modelo_generico(nome_arquivo)

@app.route('/api/historico')
def api_historico_legado():
    return jsonify([{'arquivo': f} for f in os.listdir(PROCESSED_FOLDER) if f.endswith('.docx')])

@app.route('/api/excluir/<path:filename>', methods=['DELETE'])
def excluir_arquivo(filename):
    filename = unquote(filename)
    for pasta in [UPLOAD_FOLDER, PROCESSED_FOLDER, INTERNACAO_FOLDER]:
        caminho = os.path.join(pasta, filename)
        if os.path.exists(caminho):
            os.remove(caminho)
            return "Arquivo excluído", 200
    return "Arquivo não encontrado", 404

@app.route('/api/concluir/<path:filename>', methods=['POST'])
def concluir(filename):
    filename = unquote(filename)
    origem = os.path.join(UPLOAD_FOLDER, filename)
    destino = os.path.join(PROCESSED_FOLDER, filename)
    if os.path.exists(origem):
        shutil.move(origem, destino)
        return "Concluído", 200
    return "Não encontrado", 404

@app.route('/api/baixar/<path:filename>')
def baixar(filename):
    filename = unquote(filename)
    for pasta in [UPLOAD_FOLDER, PROCESSED_FOLDER, INTERNACAO_FOLDER]:
        caminho = os.path.join(pasta, filename)
        if os.path.exists(caminho):
            return send_file(caminho, as_attachment=True)
    return "Arquivo não encontrado", 404

@app.route('/api/version')
def api_version_hemo():
    return jsonify({"version": APP_HEMO_VERSION, "filename": APP_HEMO_FILE})

@app.route('/api/download/app')
def download_hemo():
    return send_file(os.path.join(BIN_FOLDER, APP_HEMO_FILE), as_attachment=True)

@app.route('/api/recepcao/version')
def api_version_recepcao():
    return jsonify({"version": APP_RECEPCAO_VERSION, "filename": APP_RECEPCAO_FILE})

@app.route('/api/download/recepcao')
def download_recepcao():
    return send_file(os.path.join(BIN_FOLDER, APP_RECEPCAO_FILE), as_attachment=True)

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=2727, threaded=True)