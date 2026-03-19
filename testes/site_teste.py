import os
import re
import shutil
import subprocess
import tempfile
import json
from datetime import datetime
from urllib.parse import unquote

from flask import Flask, request, jsonify, render_template, send_file
from docx import Document

app = Flask(__name__)

# --- CONFIGURAÇÕES ---
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
# Usar subpasta para não misturar com o original durante testes
DATA_DIR = os.path.join(BASE_DIR, 'test_data')
UPLOAD_FOLDER = os.path.join(DATA_DIR, 'uploads')
PROCESSED_FOLDER = os.path.join(UPLOAD_FOLDER, 'concluidos')
PRE_FICHAS_FOLDER = os.path.join(DATA_DIR, 'pre_fichas')
MODELS_FOLDER = os.path.join(os.path.dirname(BASE_DIR), 'modelos')
BIN_FOLDER = os.path.join(os.path.dirname(BASE_DIR), 'bin')
MATERIAIS_FILE = os.path.join(os.path.dirname(BASE_DIR), 'materiais.json')
ATIVOS_FOLDER = os.path.join(DATA_DIR, 'procedimentos_ativos')
SALA_CONCLUIDOS_FOLDER = os.path.join(DATA_DIR, 'sala_concluidos')

EXTENSOES_PERMITIDAS = ('.doc', '.docx')

for pasta in [UPLOAD_FOLDER, PROCESSED_FOLDER, PRE_FICHAS_FOLDER, ATIVOS_FOLDER, SALA_CONCLUIDOS_FOLDER]:
    os.makedirs(pasta, exist_ok=True)

# ------------------------------------------------------------------
# FUNÇÕES AUXILIARES
# ------------------------------------------------------------------

def extrair_dados_sdt(doc):
    """
    Extrai dados de Controles de Conteúdo (Structured Document Tags) via Tags/Tags.
    """
    dados = {}
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    try:
        # Busca todos os elementos sdt
        for sdt in doc._element.xpath('.//w:sdt'):
            sdtPr = sdt.find('w:sdtPr', namespaces=ns)
            if sdtPr is None: continue
            
            tag = sdtPr.find('w:tag', namespaces=ns)
            if tag is not None:
                tag_val = tag.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                if not tag_val: continue
                
                # Busca o texto dentro do conteúdo do sdt
                content = sdt.find('w:sdtContent', namespaces=ns)
                if content is not None:
                    texto = "".join(t.text for t in content.xpath('.//w:t', namespaces=ns) if t.text)
                    if texto.strip():
                        dados[tag_val.upper()] = texto.strip()
    except Exception as e:
        print(f"Erro ao extrair via Tags: {e}")
    return dados


def extrair_dados_ficha(caminho, nome_original_fallback=None):
    """
    Tenta extrair Nome, CNS, Nascimento, Procedência e Tipo de Procedimento da ficha carregada.
    """
    dados = {
        "nome": "",
        "cns": "",
        "nasc": "",
        "procedencia": "",
        "procedimento": ""
    }
    try:
        doc = Document(caminho)
        
        # Coleta texto completo
        textos = [p.text for p in doc.paragraphs]
        for t in doc.tables:
            for r in t.rows:
                for c in r.cells:
                    textos.append(c.text)
        texto_completo = "\n".join(textos)

        # 1. Extração via Regex (Melhorado para ser mais específico e evitar capturar o rótulo)
        if not dados["nome"]:
            # Procura por Nome: [O que queremos] mas ignora o rótulo e prefixos
            match_nome = re.search(r'(?:NOME|PACIENTE|NM)\s*[:\s_.]+\s*(?:EXTERNO\s+)?([A-ZÀ-Ú]{2,}(?:\s+[A-ZÀ-Ú]{2,})+)', texto_completo, re.IGNORECASE)
            if match_nome:
                nome_limpo = match_nome.group(1).strip()
                if nome_limpo.upper() != "NOME":
                    dados["nome"] = " ".join(nome_limpo.split())

        if not dados["cns"]:
            match_cns = re.search(r'\b(\d{3}[\s\.]?\d{4}[\s\.]?\d{4}[\s\.]?\d{4})\b', texto_completo)
            if match_cns:
                dados["cns"] = re.sub(r'\D', '', match_cns.group(1))

        if not dados["nasc"]:
            # Procura data de nascimento especificamente perto da palavra nascimento
            match_nasc = re.search(r'(?:NASC|NASCIMENTO|DN)\s*[:\s_.]+\s*(\d{2}/\d{2}/\d{4})', texto_completo, re.IGNORECASE)
            if not match_nasc:
                # Se não achar com rótulo, pega a primeira data que NÃO seja a de hoje (se possível)
                hoje_str = datetime.now().strftime('%d/%m/%Y')
                todas_datas = re.findall(r'(\d{2}/\d{2}/\d{4})', texto_completo)
                for d in todas_datas:
                    if d != hoje_str:
                        dados["nasc"] = d
                        break
            else:
                dados["nasc"] = match_nasc.group(1).strip()

        if not dados["procedencia"]:
            # Melhora para pegar o valor após o rótulo, sem incluir o rótulo
            match_proc = re.search(r'(?:PROCEDÊNCIA|ORIGEM|VINDO DE|UNIDADE)\s*[:\s_.]+\s*([A-ZÀ-Ú0-9]{3,}(?:\s+[A-ZÀ-Ú0-9]{3,})*)', texto_completo, re.IGNORECASE)
            if match_proc:
                proc_limpa = match_proc.group(1).strip()
                # Remove termos genéricos que podem ter vindo junto
                proc_limpa = re.sub(r'^(?:DE ORIGEM|DA UNIDADE)\s+', '', proc_limpa, flags=re.IGNORECASE)
                dados["procedencia"] = proc_limpa

        # 2. Extração via Tags (Fallback se o Regex falhar)
        dados_sdt = extrair_dados_sdt(doc)
        if not dados["nome"] and "NOME" in dados_sdt: dados["nome"] = dados_sdt["NOME"].strip()
        if not dados["cns"] and "CNS" in dados_sdt: dados["cns"] = re.sub(r'\D', '', dados_sdt["CNS"])
        if not dados["nasc"] and "NASC" in dados_sdt: dados["nasc"] = dados_sdt["NASC"]
        if not dados["procedencia"] and "PROCEDENCIA" in dados_sdt: dados["procedencia"] = dados_sdt["PROCEDENCIA"].strip()

        # 3. Fallback do Nome do Arquivo (USAR O NOME ORIGINAL ENVIADO)
        if not dados["nome"] or dados["nome"].upper() == "EXTERNO":
            ref_name = nome_original_fallback if nome_original_fallback else os.path.basename(caminho)
            nome_arq_limpo = re.sub(r'^\d{8}_', '', ref_name).rsplit('.', 1)[0]
            nome_arq_limpo = re.sub(r'_(CATETERISMO|ANGIOPLASTIA|PTCA|CINE)$', '', nome_arq_limpo, flags=re.IGNORECASE)
            dados["nome"] = nome_arq_limpo.replace('_', ' ').replace('-', ' ').strip()

        if not dados["nome"] or len(dados["nome"]) < 3:
            dados["nome"] = "Externo"

        return dados
    except Exception as e:
        print(f"Erro na extração: {e}")
        return dados


def gerar_laudo_individual(tipo_label, nome_paciente, dados_extracao, materiais=None):
    """
    Gera um laudo específico (CAT ou ANGIO).
    """
    modelos = {
        'CATETERISMO': 'Laudo de cateterismo.docx',
        'ANGIOPLASTIA': 'Laudo de Angioplastia.docx'
    }
    
    modelo_nome = modelos.get(tipo_label)
    if not modelo_nome: return None
    
    caminho_modelo = os.path.join(MODELS_FOLDER, modelo_nome)
    if not os.path.exists(caminho_modelo):
        print(f"Modelo não encontrado: {caminho_modelo}")
        return None

    try:
        txt_materiais = ""
        if materiais:
            # Limpa espaços excessivos entre palavras em cada material
            materiais_limpos = [" ".join(m.split()) for m in materiais if m and "contraste" not in m.lower()]
            txt_materiais = "\n".join([f"- {m}" for m in materiais_limpos])

        # Garante que não fiquem vazios para evitar manter o placeholder no documento
        procedencia = dados_extracao.get("procedencia", "")
        if not procedencia: procedencia = "Não Identificada"

        dados_laudo = {
            "{{NOME}}": to_pascal_case(nome_paciente),
            "{{CNS}}": dados_extracao.get("cns", ""),
            "{{NASC}}": dados_extracao.get("nasc", ""),
            "{{NASCIMENTO}}": dados_extracao.get("nasc", ""),
            "{{PROCEDENCIA}}": to_pascal_case(procedencia),
            "{{DATA_HOJE}}": datetime.now().strftime('%d/%m/%Y'),
            "{{NUM_EXAME}}": "{{NUM_EXAME}}", # Mantém para o Java
            "{{MATERIAIS}}": txt_materiais,
            "{{LOGRADOURO}}": "", "{{BAIRRO}}": "", "{{CIDADE}}": "", "{{UF}}": "",
            "{{RG}}": "", "{{CPF}}": "", "{{MAE}}": "", "{{PAI}}": "",
            "{{MATRICULA}}": "", "{{CHAVE}}": "", "{{TELEFONES}}": ""
        }

        doc = Document(caminho_modelo)
        substituir_placeholders_py(doc, dados_laudo)

        data_str = datetime.now().strftime('%Y%m%d')
        nome_limpo = re.sub(r'[\\/*?:"<>|]', "", nome_paciente.replace(' ', '_'))
        nome_arquivo = f"{data_str}_{nome_limpo}_{tipo_label}.docx"

        caminho_final = os.path.join(UPLOAD_FOLDER, nome_arquivo)
        contador = 1
        while os.path.exists(caminho_final):
            caminho_final = os.path.join(UPLOAD_FOLDER, f"{data_str}_{nome_limpo}_{tipo_label}_{contador}.docx")
            contador += 1

        doc.save(caminho_final)
        return os.path.basename(caminho_final)
    except Exception as e:
        print(f"Erro ao gerar {tipo_label}: {e}")
        return None


def gerar_laudos_automaticos(nome_paciente, dados_extracao):
    """
    Gera o laudo correspondente ao procedimento detectado na ficha.
    Se não for identificado, gera Cateterismo por padrão.
    """
    procedimento = dados_extracao.get("procedimento", "CATETERISMO")
    
    if "ANGIOPLASTIA" in procedimento:
        angio = gerar_laudo_individual('ANGIOPLASTIA', nome_paciente, dados_extracao)
        return [angio] if angio else []
    else:
        cat = gerar_laudo_individual('CATETERISMO', nome_paciente, dados_extracao)
        return [cat] if cat else []


def extrair_nome(caminho):
    dados = extrair_dados_ficha(caminho)
    return dados["nome"]


def gerar_nome_seguro(nome_paciente, pasta_destino):
    data = datetime.now().strftime('%Y%m%d')
    nome_base = f"{data}_{nome_paciente.replace(' ', '_')}"

    nome_base = re.sub(r'[\\/*?:"<>|]', "", nome_base)
    caminho = os.path.join(pasta_destino, f"{nome_base}.docx")
    contador = 1

    while os.path.exists(caminho):
        caminho = os.path.join(pasta_destino, f"{nome_base}_{contador}.docx")
        contador += 1

    return os.path.basename(caminho)


def converter_doc_para_docx(caminho_doc):
    pasta = os.path.dirname(caminho_doc)
    try:
        subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "docx", caminho_doc, "--outdir", pasta],
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
            timeout=30,
            check=True
        )
        novo = caminho_doc.rsplit('.', 1)[0] + '.docx'
        return novo if os.path.exists(novo) else None
    except Exception:
        return None


def to_pascal_case(valor):
    if not valor:
        return ""
    partes = re.split(r"\s+", valor.strip())
    return " ".join(p[:1].upper() + p[1:].lower() for p in partes if p)


def substituir_placeholders_py(doc, dados):
    """
    Percorre paragrafos e tabelas do documento e substitui os placeholders.
    Lida com o fato de que o Word as vezes quebra o placeholder em varios 'runs'.
    """
    def process_paragraph(p):
        for key, value in dados.items():
            if key in p.text:
                # Tenta substituir primeiro em cada run individualmente
                for run in p.runs:
                    if key in run.text:
                        run.text = run.text.replace(key, str(value))
                
                # Se ainda existir a chave no texto total (estava dividida entre runs)
                if key in p.text:
                    # Estrategia de fallback: reconstrói o texto no primeiro run e limpa os outros
                    # Isso evita o problema de "espaços grandes" ou falha na substituição
                    new_text = p.text.replace(key, str(value))
                    if p.runs:
                        p.runs[0].text = new_text
                        for i in range(1, len(p.runs)):
                            p.runs[i].text = ""

    for p in doc.paragraphs:
        process_paragraph(p)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    process_paragraph(p)

# ------------------------------------------------------------------
# ROTAS WEB
# ------------------------------------------------------------------

@app.route('/', methods=['GET', 'POST'])
def index():
    enviados, erros = [], []

    if request.method == 'POST':
        arquivos = request.files.getlist('files')

        for f in arquivos:
            if not f.filename.lower().endswith(EXTENSOES_PERMITIDAS):
                erros.append(f"{f.filename}: Extensão não permitida")
                continue

            caminho_original = None
            caminho_docx = None

            try:
                sufixo = '.doc' if f.filename.lower().endswith('.doc') else '.docx'
                with tempfile.NamedTemporaryFile(delete=False, suffix=sufixo) as tmp:
                    f.save(tmp.name)
                    caminho_original = tmp.name

                caminho_docx = caminho_original

                if f.filename.lower().endswith('.doc'):
                    caminho_docx = converter_doc_para_docx(caminho_original)
                    if not caminho_docx:
                        erros.append(f"{f.filename}: Falha ao converter .doc para .docx")
                        continue

                dados_ficha = extrair_dados_ficha(caminho_docx, f.filename)
                nome_paciente = dados_ficha["nome"]
                nome_final = gerar_nome_seguro(nome_paciente, UPLOAD_FOLDER)
                caminho_destino = os.path.join(UPLOAD_FOLDER, nome_final)

                shutil.move(caminho_docx, caminho_destino)
                
                msg = f"{f.filename} → {nome_final}"
                enviados.append(msg)

            except Exception as e:
                erros.append(f"{f.filename}: {str(e)}")

            finally:
                for c in [caminho_original, caminho_docx]:
                    try:
                        if c and os.path.exists(c) and not c.startswith(UPLOAD_FOLDER):
                            os.remove(c)
                    except FileNotFoundError:
                        pass

    return render_template('index.html', enviados=enviados, erros=erros)


@app.route('/historico')
def historico():
    pendentes, concluidos, em_processamento = [], [], []

    for f in sorted(os.listdir(PROCESSED_FOLDER), reverse=True):
        if f.endswith('.docx'):
            concluidos.append(f)

    for f in sorted(os.listdir(UPLOAD_FOLDER), reverse=True):
        if f.endswith('.docx'):
            caminho = os.path.join(UPLOAD_FOLDER, f)
            stat = os.stat(caminho)
            if stat.st_mtime > stat.st_ctime + 1:
                em_processamento.append(f)
            else:
                pendentes.append(f)

    return render_template(
        'historico.html',
        pendentes=pendentes,
        concluidos=concluidos,
        em_processamento=em_processamento
    )


@app.route('/view/<path:filename>')
def view_file(filename):
    filename = unquote(filename)
    for pasta in [UPLOAD_FOLDER, PROCESSED_FOLDER]:
        caminho = os.path.join(pasta, filename)
        if os.path.exists(caminho):
            doc = Document(caminho)
            texto = []

            for p in doc.paragraphs:
                texto.append(p.text)

            for t in doc.tables:
                for r in t.rows:
                    texto.append("\t".join(c.text for c in r.cells))

            return render_template('view.html', filename=filename, content="\n".join(texto))

    return "Arquivo não encontrado", 404

# ------------------------------------------------------------------
# ROTAS API
# ------------------------------------------------------------------

# --- FICHAS PRÉ-PRONTAS (RASCUNHOS) ---

@app.route('/api/fichas/pre', methods=['POST'])
def salvar_rascunho():
    dados = request.get_json()
    if not dados or 'nome_rascunho' not in dados:
        return jsonify({"erro": "Nome do rascunho é obrigatório"}), 400

    nome_rascunho = re.sub(r'[\\/*?:"<>|]', "", dados['nome_rascunho'])
    data_hoje = datetime.now().strftime('%Y-%m-%d')
    nome_arquivo = f"{data_hoje}_{nome_rascunho}.json"

    caminho_arquivo = os.path.join(PRE_FICHAS_FOLDER, nome_arquivo)

    # Evita sobrescrever
    contador = 1
    while os.path.exists(caminho_arquivo):
        caminho_arquivo = os.path.join(PRE_FICHAS_FOLDER, f"{data_hoje}_{nome_rascunho}_{contador}.json")
        contador += 1

    try:
        with open(caminho_arquivo, 'w', encoding='utf-8') as f:
            json.dump(dados, f, ensure_ascii=False, indent=4)
        return jsonify({"sucesso": "Rascunho salvo com sucesso", "filename": os.path.basename(caminho_arquivo)}), 201
    except Exception as e:
        return jsonify({"erro": str(e)}), 500

@app.route('/api/fichas/pre', methods=['GET'])
def listar_rascunhos():
    try:
        # Lista e ordena por data (mais recentes primeiro)
        arquivos = sorted(os.listdir(PRE_FICHAS_FOLDER), reverse=True)
        return jsonify([f for f in arquivos if f.endswith('.json')])
    except Exception as e:
        return jsonify({"erro": str(e)}), 500

@app.route('/api/fichas/pre/<path:filename>', methods=['GET'])
def carregar_rascunho(filename):
    caminho = os.path.join(PRE_FICHAS_FOLDER, filename)
    if not os.path.exists(caminho):
        return jsonify({"erro": "Rascunho não encontrado"}), 404
    return send_file(caminho, mimetype='application/json')

@app.route('/api/fichas/pre/<path:filename>', methods=['DELETE'])
def deletar_rascunho(filename):
    caminho = os.path.join(PRE_FICHAS_FOLDER, filename)
    if not os.path.exists(caminho):
        return jsonify({"erro": "Rascunho não encontrado"}), 404
    try:
        os.remove(caminho)
        return jsonify({"sucesso": "Rascunho deletado"}), 200
    except Exception as e:
        return jsonify({"erro": str(e)}), 500


@app.route('/api/laudos/internado', methods=['POST'])
def gerar_laudo_internado():
    """
    Recebe dados de um paciente internado e gera AMBOS os laudos (CAT e ANGIO)
    na pasta 'uploads', deixando-os prontos para o Java preencher o número.
    """
    dados_req = request.get_json()
    if not dados_req:
        return jsonify({"erro": "Requisição sem JSON"}), 400

    campos_necessarios = ["nome", "cns", "nasc", "procedencia"]
    if not all(campo in dados_req for campo in campos_necessarios):
        return jsonify({"erro": "Dados incompletos no JSON"}), 400

    try:
        nome_paciente = dados_req.get("nome", "")
        # Reutiliza a função de geração automática
        gerados = gerar_laudos_automaticos(nome_paciente, dados_req)

        if not gerados:
            return jsonify({"erro": "Nenhum laudo foi gerado. Verifique os modelos."}), 500

        return jsonify({
            "sucesso": "Laudos gerados com sucesso",
            "filenames": gerados
        }), 201

    except Exception as e:
        return jsonify({"erro": f"Falha ao gerar documentos: {str(e)}"}), 500


@app.route('/api/pendentes')
def api_pendentes():
    # Usamos um dicionário para evitar duplicados, a chave é o nome do arquivo ou nome do paciente
    # Prioridade: ATIVOS > CONCLUIDOS > UPLOAD
    pendentes_dict = {}
    client = request.args.get('client')
    
    # 1. Pacientes Ativos na Sala (JSONs em procedimentos_ativos)
    for f in os.listdir(ATIVOS_FOLDER):
        if f.endswith('.json'):
            try:
                with open(os.path.join(ATIVOS_FOLDER, f), 'r', encoding='utf-8') as file:
                    dados = json.load(file)
                    nome_pac = to_pascal_case(dados.get("nome", f))
                    pendentes_dict[f] = {
                        "arquivo": f,
                        "nome": nome_pac,
                        "status": "EM SALA",
                        "origem": "TABLET",
                        "procedencia": dados.get("procedencia", ""),
                        "nasc": dados.get("nasc", ""),
                        "cns": dados.get("cns", ""),
                        "tipo_procedimento": "CAT + ANGIO" if dados.get("evoluiu_angioplastia") else "CATETERISMO",
                        "materiais_cat": "\n".join(dados.get("materiais_cateterismo", [])),
                        "materiais_angio": "\n".join(dados.get("materiais_angioplastia", []))
                    }
            except Exception as e: print(f"Erro lendo ativo {f}: {e}")

    # 2. Pacientes que Finalizaram na Sala (JSONs em sala_concluidos)
    for f in os.listdir(SALA_CONCLUIDOS_FOLDER):
        if f.endswith('.json') and f not in pendentes_dict:
            try:
                with open(os.path.join(SALA_CONCLUIDOS_FOLDER, f), 'r', encoding='utf-8') as file:
                    dados = json.load(file)
                    nome_pac = to_pascal_case(dados.get("nome", f))
                    pendentes_dict[f] = {
                        "arquivo": f,
                        "nome": nome_pac,
                        "status": "AGUARDANDO LAUDO",
                        "origem": "TABLET (CONCLUÍDO)",
                        "procedencia": dados.get("procedencia", ""),
                        "nasc": dados.get("nasc", ""),
                        "cns": dados.get("cns", ""),
                        "tipo_procedimento": "CAT + ANGIO" if dados.get("evoluiu_angioplastia") else "CATETERISMO",
                        "materiais_cat": "\n".join(dados.get("materiais_cateterismo", [])),
                        "materiais_angio": "\n".join(dados.get("materiais_angioplastia", []))
                    }
            except Exception as e: print(f"Erro lendo concluído {f}: {e}")

    # 3. Fichas Manuais (Word em uploads)
    if client != 'java':
        for f in os.listdir(UPLOAD_FOLDER):
            if f.endswith('.docx') and not f.startswith('~$') and not f.startswith('INT_'):
                # Evita duplicar se o JSON do tablet já existe com nome similar
                ja_existe = any(f.replace('.docx', '') in k for k in pendentes_dict.keys())
                if not ja_existe:
                    caminho = os.path.join(UPLOAD_FOLDER, f)
                    dados = extrair_dados_ficha(caminho, f)
                    pendentes_dict[f] = {
                        "arquivo": f,
                        "nome": to_pascal_case(dados["nome"]),
                        "status": "MANUAL",
                        "origem": "UPLOAD",
                        "procedencia": dados["procedencia"],
                        "nasc": dados["nasc"],
                        "cns": dados["cns"],
                        "tipo_procedimento": dados.get("procedimento", "CATETERISMO"),
                        "materiais_cat": "",
                        "materiais_angio": ""
                    }
            
    return jsonify(list(pendentes_dict.values()))


@app.route('/api/baixar/json/<path:filename>')
def api_baixar_json(filename):
    filename = unquote(filename)
    for pasta in [ATIVOS_FOLDER, SALA_CONCLUIDOS_FOLDER]:
        caminho = os.path.join(pasta, filename)
        if os.path.exists(caminho):
            return send_file(caminho, mimetype='application/json')
    return "Não encontrado", 404


@app.route('/api/internacao/listar')
def api_listar_internacao():
    lista = []
    for f in os.listdir(ATIVOS_FOLDER):
        if f.endswith('.json'):
            with open(os.path.join(ATIVOS_FOLDER, f), 'r', encoding='utf-8') as file:
                dados = json.load(file)
                lista.append({
                    "nome": dados.get("nome", "Desconhecido"),
                    "data": dados.get("inicio", ""),
                    "arquivo": f,
                    "procedencia": dados.get("procedencia", "Não informada")
                })
    return jsonify(lista)


@app.route('/api/internacao/gerar', methods=['POST'])
def gerar_docs_internacao():
    dados_req = request.get_json()
    nome_paciente = dados_req.get('nome', 'Paciente')
    modelos_pedidos = dados_req.get('modelos', [])
    
    # Mapeamento de nomes de modelos do Java para arquivos reais
    mapeamento = {
        "solicitacao_angio_sus.docx": "pedido_angio_sus.docx",
        "justificativa_angio.docx": "justificativa_angioplastia.docx",
        "evolucao_tasy.docx": "internacao.docx" # Fallback se faltar
    }
    
    arquivos_gerados = []
    data_hoje = datetime.now().strftime('%d/%m/%Y')
    
    for mod_nome in modelos_pedidos:
        real_mod = mapeamento.get(mod_nome, mod_nome)
        caminho_modelo = os.path.join(MODELS_FOLDER, real_mod)
        
        if not os.path.exists(caminho_modelo):
            print(f"Modelo não encontrado: {caminho_modelo}")
            continue
            
        try:
            doc = Document(caminho_modelo)
            
            # Dados para preenchimento
            dados_preencher = {
                "{{NOME}}": to_pascal_case(nome_paciente),
                "{{DATA_HOJE}}": data_hoje,
                "{{ARTERIAS}}": dados_req.get('arterias', ''),
                "{{STENTS}}": dados_req.get('stents', ''),
                "{{PROCEDENCIA}}": to_pascal_case(dados_req.get('procedencia', '')),
                "{{OBS_TXT}}": dados_req.get('obs_txt', ''),
                "{{CHK_CLINICO}}": dados_req.get('chk_clinico', ' '),
                "{{CHK_ANGIO}}": dados_req.get('chk_angio', ' '),
                "{{CHK_CIRURGIA}}": dados_req.get('chk_cirurgia', ' ')
            }
            
            substituir_placeholders_py(doc, dados_preencher)
            
            # Salva no UPLOAD_FOLDER para permitir download posterior
            nome_limpo = re.sub(r'[\\/*?:"<>|]', "", nome_paciente.replace(' ', '_'))
            ext = real_mod.split('.')[-1]
            nome_final = f"INT_{nome_limpo}_{real_mod.split('.')[0]}_{datetime.now().strftime('%H%M%S')}.{ext}"
            caminho_final = os.path.join(UPLOAD_FOLDER, nome_final)
            
            doc.save(caminho_final)
            arquivos_gerados.append(nome_final)
            
        except Exception as e:
            print(f"Erro ao gerar {mod_nome}: {e}")
            
    return jsonify({"arquivos_gerados": arquivos_gerados}), 201


@app.route('/api/historico')
def api_historico():
    lista = []
    for f in os.listdir(PROCESSED_FOLDER):
        if f.endswith('.docx'):
            lista.append({"arquivo": f})
    return jsonify(lista)


@app.route('/api/excluir/<filename>', methods=['DELETE'])
def excluir_arquivo(filename):
    filename = unquote(filename)
    deletado = False

    for pasta in [UPLOAD_FOLDER, PROCESSED_FOLDER]:
        caminho = os.path.join(pasta, filename)
        if os.path.exists(caminho):
            os.remove(caminho)
            deletado = True

    return ("Arquivo excluído", 200) if deletado else ("Arquivo não encontrado", 404)


@app.route('/api/concluir/<filename>', methods=['POST'])
def concluir(filename):
    filename = unquote(filename)
    origem = os.path.join(UPLOAD_FOLDER, filename)
    destino = os.path.join(PROCESSED_FOLDER, filename)

    if os.path.exists(origem):
        shutil.move(origem, destino)
        return "Concluído", 200

    return "Não encontrado", 404


@app.route('/api/baixar/<filename>')
def baixar(filename):
    filename = unquote(filename)

    for pasta in [UPLOAD_FOLDER, PROCESSED_FOLDER]:
        caminho = os.path.join(pasta, filename)
        if os.path.exists(caminho):
            return send_file(caminho, as_attachment=True)

    return "Arquivo não encontrado", 404


@app.route('/api/modelos/<path:filename>')
def serve_modelos(filename):
    filename = unquote(filename)
    caminho = os.path.join(MODELS_FOLDER, filename)
    if os.path.exists(caminho):
        return send_file(caminho, as_attachment=True)
    return f"Modelo não encontrado: {filename}", 404


@app.route('/api/modelo/<tipo>')
def modelo(tipo):
    nome = {
        'cateterismo': 'Laudo de cateterismo.docx',
        'angioplastia': 'Laudo de Angioplastia.docx'
    }.get(tipo)

    if not nome:
        return "Tipo inválido", 400

    caminho = os.path.join(MODELS_FOLDER, nome)
    return send_file(caminho, as_attachment=True) if os.path.exists(caminho) else ("Modelo não encontrado", 404)


# --- ROTAS DA SALA (TABLET) ---

@app.route('/tablet')
def tablet_index():
    return render_template('tablet.html')

@app.route('/api/materiais', methods=['GET'])
def get_materiais():
    if not os.path.exists(MATERIAIS_FILE):
        return jsonify({"configuracoes": {}, "catalogo": {}})
    with open(MATERIAIS_FILE, 'r', encoding='utf-8') as f:
        materiais_data = json.load(f)
    
    procedimento = request.args.get('procedimento')
    if procedimento == 'cateterismo':
        # Filtra apenas o que é pertinente ao cateterismo
        materiais_data['catalogo'] = {
            "acesso_e_diagnostico": materiais_data['catalogo'].get('acesso_e_diagnostico', [])
        }
    return jsonify(materiais_data)

@app.route('/api/materiais', methods=['POST'])
def save_material():
    """Adiciona ou atualiza um material no catálogo visualmente."""
    novo_item = request.get_json()
    tipo_cat = novo_item.get('categoria') 
    
    if not os.path.exists(MATERIAIS_FILE):
        materiais_data = {
            "configuracoes": {"kit_padrao_cateterismo": []},
            "catalogo": {
                "acesso_e_diagnostico": [],
                "intervencao": {"stents": [], "baloes": [], "guias": []}
            }
        }
    else:
        with open(MATERIAIS_FILE, 'r', encoding='utf-8') as f:
            materiais_data = json.load(f)

    catalogo = materiais_data.get('catalogo', {})
    intervencao = catalogo.get('intervencao', {})

    if tipo_cat == 'acesso_e_diagnostico':
        catalogo['acesso_e_diagnostico'].append(novo_item)
    elif tipo_cat in ['stents', 'baloes', 'guias']:
        if tipo_cat in ['stents', 'baloes']:
            marca = novo_item.get('marca')
            found = False
            for item in intervencao.get(tipo_cat, []):
                if item['marca'] == marca and item.get('tipo') == novo_item.get('tipo'):
                    item['medidas'] = list(set(item.get('medidas', []) + novo_item.get('medidas', [])))
                    found = True
                    break
            if not found:
                intervencao.setdefault(tipo_cat, []).append(novo_item)
        else:
            intervencao.setdefault(tipo_cat, []).append(novo_item)
    else:
        return jsonify({"erro": "Categoria inválida"}), 400

    with open(MATERIAIS_FILE, 'w', encoding='utf-8') as f:
        json.dump(materiais_data, f, indent=4, ensure_ascii=False)
    return jsonify({"status": "sucesso"}), 201

@app.route('/api/sala/ativo', methods=['GET'])
def get_paciente_em_mesa():
    """Retorna o paciente que está atualmente em procedimento, se houver."""
    arquivos = os.listdir(ATIVOS_FOLDER)
    if not arquivos:
        return jsonify({"status": "vazio"})
    
    with open(os.path.join(ATIVOS_FOLDER, arquivos[0]), 'r', encoding='utf-8') as f:
        return jsonify(json.load(f))

@app.route('/api/sala/iniciar', methods=['POST'])
def iniciar_mesa():
    """Inicia um procedimento para um paciente selecionado."""
    dados = request.get_json()
    nome_paciente = dados.get('nome')
    procedimento = dados.get('procedimento', 'cateterismo')
    
    # Informações extras extraídas para o laudo
    cns = dados.get('cns', '')
    nasc = dados.get('nasc', '')
    procedencia = dados.get('procedencia', '')
    
    # Limpa procedimentos ativos anteriores
    for f in os.listdir(ATIVOS_FOLDER):
        os.remove(os.path.join(ATIVOS_FOLDER, f))

    info = {
        "nome": nome_paciente,
        "cns": cns,
        "nasc": nasc,
        "procedencia": procedencia,
        "procedimento": procedimento,
        "inicio": datetime.now().strftime('%H:%M:%S'),
        "materiais_cateterismo": [],
        "materiais_angioplastia": [],
        "evoluiu_angioplastia": False
    }

    # Se for Cateterismo, carrega o Kit Padrão na lista de cateterismo
    if procedimento == 'cateterismo':
        if os.path.exists(MATERIAIS_FILE):
            with open(MATERIAIS_FILE, 'r', encoding='utf-8') as f:
                materiais_data = json.load(f)
                kit = materiais_data.get('configuracoes', {}).get('kit_padrao_cateterismo', [])
                for item in kit:
                    nome_item = item['nome']
                    if 'calibre_padrao' in item:
                        nome_item += f" {item['calibre_padrao']}"
                    info['materiais_cateterismo'].append(nome_item)
    
    filename = f"ativo_{re.sub(r'[^a-zA-Z0-9]', '_', nome_paciente)}.json"
    with open(os.path.join(ATIVOS_FOLDER, filename), 'w', encoding='utf-8') as f:
        json.dump(info, f, indent=4)
        
    return jsonify({"status": "iniciado"}), 201

@app.route('/api/sala/atualizar', methods=['POST'])
def atualizar_mesa():
    """Adiciona material ou muda status (CAT -> ANGIO)."""
    dados = request.get_json()
    arquivos = os.listdir(ATIVOS_FOLDER)
    if not arquivos:
        return jsonify({"erro": "Nenhum paciente em mesa"}), 404

    caminho = os.path.join(ATIVOS_FOLDER, arquivos[0])
    with open(caminho, 'r', encoding='utf-8') as f:
        info = json.load(f)

    if 'material' in dados:
        material_nome = dados['material']
        calibre = dados.get('calibre')
        if calibre:
            material_nome = f"{material_nome} {calibre}"
        
        # Define em qual lista o material deve entrar
        if info.get('evoluiu_angioplastia'):
            info['materiais_angioplastia'].append(material_nome)
        else:
            info['materiais_cateterismo'].append(material_nome)
        
    if 'evoluiu_angioplastia' in dados:
        foi_para_angio = dados['evoluiu_angioplastia']
        # Se mudou de False para True, gera o laudo de angioplastia
        if foi_para_angio and not info.get('evoluiu_angioplastia'):
            info['evoluiu_angioplastia'] = True
            
            # Adiciona materiais automáticos solicitados: insuflador e fio guia
            # Evita duplicar se por algum motivo já estiverem lá
            materiais_auto = ["insuflador 20atm", "fio guia 0.014x180cm"]
            for m_auto in materiais_auto:
                if m_auto not in info['materiais_angioplastia']:
                    info['materiais_angioplastia'].append(m_auto)
            
            # GERA O LAUDO DE ANGIOPLASTIA AQUI
            dados_extracao = {
                "cns": info.get("cns", ""),
                "nasc": info.get("nasc", ""),
                "procedencia": info.get("procedencia", "")
            }
            # Combina materiais (já limpa contraste na função de geração)
            todos_materiais = info['materiais_cateterismo'] + info['materiais_angioplastia']
            
            gerar_laudo_individual('ANGIOPLASTIA', info['nome'], dados_extracao, todos_materiais)

    with open(caminho, 'w', encoding='utf-8') as f:
        json.dump(info, f, indent=4)
        
    return jsonify({"status": "atualizado"})

@app.route('/api/finalizar_laudo/<filename>', methods=['POST'])
def finalizar_laudo_tablet(filename):
    """Remove o JSON de sala_concluidos ou ativos quando o laudo for gerado no Java."""
    filename = unquote(filename)
    deletado = False
    for pasta in [SALA_CONCLUIDOS_FOLDER, ATIVOS_FOLDER]:
        caminho = os.path.join(pasta, filename)
        if os.path.exists(caminho):
            os.remove(caminho)
            deletado = True
    return ("OK", 200) if deletado else ("Não encontrado", 404)


@app.route('/api/sala/finalizar', methods=['POST'])
def finalizar_mesa():
    """Move o paciente da sala ativa para a fila de laudos pendentes (sala_concluidos)."""
    arquivos = os.listdir(ATIVOS_FOLDER)
    if not arquivos:
        return jsonify({"erro": "Nenhum paciente em mesa"}), 404

    # Pega o primeiro (e único) paciente em mesa
    nome_arq = arquivos[0]
    origem = os.path.join(ATIVOS_FOLDER, nome_arq)
    destino = os.path.join(SALA_CONCLUIDOS_FOLDER, nome_arq)
    
    try:
        shutil.move(origem, destino)
        return jsonify({"status": "finalizado", "arquivo": nome_arq})
    except Exception as e:
        return jsonify({"erro": str(e)}), 500

# ------------------------------------------------------------------

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=2400, threaded=True)