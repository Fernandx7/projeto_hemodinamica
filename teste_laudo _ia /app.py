import os
import time
import shutil
from fastapi import FastAPI, Request, HTTPException, File, UploadFile, Form
from fastapi.responses import HTMLResponse, FileResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from docx import Document
import google.generativeai as genai
from pydub import AudioSegment
from typing import Optional

app = FastAPI()

# Configuration
CHAVE_API_PATH = "chaveapi"
TEMPLATE_DOCX = "Laudo de cateterismo.docx"
RESULTS_DIR = "results"
TEMP_DIR = "temp_uploads"

if not os.path.exists(RESULTS_DIR):
    os.makedirs(RESULTS_DIR)
if not os.path.exists(TEMP_DIR):
    os.makedirs(TEMP_DIR)

# Load API Key
try:
    with open(CHAVE_API_PATH, "r") as f:
        api_key = f.read().strip()
    genai.configure(api_key=api_key)
except Exception as e:
    print(f"Error loading API key: {e}")

# Static and Templates
app.mount("/static", StaticFiles(directory="static"), name="static")
templates = Jinja2Templates(directory="templates")

def get_template_text(file_path):
    try:
        doc = Document(file_path)
        return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    except Exception as e:
        print(f"Error reading DOCX: {e}")
        return ""

@app.get("/", response_class=HTMLResponse)
async def read_root(request: Request):
    return templates.TemplateResponse("index.html", {"request": request, "template_name": TEMPLATE_DOCX})

@app.post("/process")
async def process_report(
    audio_file: UploadFile = File(...),
    template_name: str = Form("Laudo de cateterismo.docx")
):
    temp_file_path = os.path.join(TEMP_DIR, audio_file.filename)
    mp3_file_path = os.path.join(TEMP_DIR, f"conv_{int(time.time())}.mp3")
    
    try:
        # 1. Save uploaded file temporarily
        with open(temp_file_path, "wb") as buffer:
            shutil.copyfileobj(audio_file.file, buffer)

        # 2. Convert to MP3 to ensure Gemini compatibility
        print(f"Converting {audio_file.filename} to MP3...")
        try:
            audio = AudioSegment.from_file(temp_file_path)
            audio.export(mp3_file_path, format="mp3")
            process_path = mp3_file_path
        except Exception as e:
            print(f"Conversion failed, using original: {e}")
            process_path = temp_file_path

        # 3. Prepare Gemini model
        model = genai.GenerativeModel("models/gemini-2.5-flash")
        
        # 4. Upload audio file to Gemini
        print(f"Uploading to Gemini: {process_path}")
        sample_file = genai.upload_file(path=process_path, display_name=f"Audio_{audio_file.filename}")
        
        # 5. Wait for processing
        while sample_file.state.name == "PROCESSING":
            print(f"Processing audio: {sample_file.name} (State: {sample_file.state.name})...")
            time.sleep(2)
            sample_file = genai.get_file(sample_file.name)
            
        if sample_file.state.name == "FAILED":
            error_msg = f"Audio processing failed on Gemini side. State: {sample_file.state.name}"
            if hasattr(sample_file, 'error') and sample_file.error:
                error_msg += f" Details: {sample_file.error}"
            print(f"ERROR: {error_msg}")
            raise HTTPException(status_code=500, detail=error_msg)

        # 6. Get template text
        template_text = get_template_text(template_name)
        
        # 7. Create specialized prompt based on patterns
        if "Angioplastia" in template_name:
            instructions = """
            Este é um laudo de ANGIOPLASTIA CORONARIANA.
            Certifique-se de incluir detalhes sobre:
            - Angiografia Pré-Procedimento (identificando as lesões).
            - Detalhes do Procedimento: tipo de punção, introdutor, cateter-guia, fio guia, stent (modelo, medidas em mm), pressão de liberação (ATM) e fluxo final (ex: TIMI III).
            - Conclusão: Sucesso do procedimento com implante de stent.
            """
        else:
            instructions = """
            Este é um laudo de CATETERISMO (CORONARIOGRAFIA E VENTRICULOGRAFIA ESQUERDA).
            Certifique-se de preencher as seções:
            - CORONARIOGRAFIA: Tronco, Descendente Anterior, Circunflexa e Coronária Direita (detalhando se há obstruções ou se estão livres).
            - VENTRICULOGRAFIA ESQUERDA: Mobilidade parietal, valva mitral e gradiente.
            - CONCLUSÃO: Resumo dos achados obstrutivos e da função global do ventrículo esquerdo.
            """

        prompt = f"""
        Você é um médico especialista em cardiologia intervencionista. 
        Ouça o áudio anexo com atenção e preencha o modelo de laudo fornecido.
        
        {instructions}
        
        Baseie-se estritamente nas informações ditas no áudio. Use terminologia médica formal.
        Retorne APENAS o texto completo do laudo preenchido, mantendo rigorosamente a estrutura original do modelo abaixo.
        
        MODELO:
        ---
        {template_text}
        ---
        """
        
        # 8. Generate content
        print(f"Generating report content for {template_name}...")
        response = model.generate_content([prompt, sample_file])
        filled_text = response.text
        
        # 9. Create new DOCX
        output_prefix = "Laudo_Angio" if "Angioplastia" in template_name else "Laudo_Cat"
        output_filename = f"{output_prefix}_{os.path.splitext(audio_file.filename)[0]}_{int(time.time())}.docx"
        output_path = os.path.join(RESULTS_DIR, output_filename)
        
        new_doc = Document()
        for line in filled_text.split('\n'):
            if line.strip():
                new_doc.add_paragraph(line)
        
        new_doc.save(output_path)
        
        # Clean up Gemini and local temp files
        genai.delete_file(sample_file.name)
        if os.path.exists(temp_file_path): os.remove(temp_file_path)
        if os.path.exists(mp3_file_path): os.remove(mp3_file_path)
        
        return {"status": "success", "filename": output_filename, "download_url": f"/download/{output_filename}"}
        
    except Exception as e:
        print(f"Error in processing: {e}")
        if os.path.exists(temp_file_path): os.remove(temp_file_path)
        if os.path.exists(mp3_file_path): os.remove(mp3_file_path)
        return {"status": "error", "message": str(e)}

@app.get("/download/{filename}")
async def download_file(filename: str):
    file_path = os.path.join(RESULTS_DIR, filename)
    if os.path.exists(file_path):
        return FileResponse(file_path, filename=filename, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    raise HTTPException(status_code=404, detail="File not found")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
