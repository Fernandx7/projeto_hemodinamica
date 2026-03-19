import os
import re
import shutil
import subprocess
from datetime import datetime
from urllib.parse import unquote
from flask import Flask, request, jsonify, render_template, send_file
from docx import Document
import tempfile
import json

app = Flask(__name__)

# --- CONFIGURAÇÕES ---
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
UPLOAD_FOLDER = os.path.join(BASE_DIR, 'uploads')
PROCESSED_FOLDER = os.path.join(UPLOAD_FOLDER, 'concluidos')
PRE_FICHAS_FOLDER = os.path.join(BASE_DIR, 'pre_fichas')
MODELS_FOLDER = os.path.join(BASE_DIR, 'modelos')
BIN_FOLDER = os.path.join(BASE_DIR, 'bin')

# --- CONTROLE DE VERSÃO ---
APP_HEMO_VERSION = "9.0.3"
APP_HEMO_FILE = "AppHemo.jar"
APP_RECEPCAO_VERSION = "5.1"
APP_RECEPCAO_FILE = "AppRecepcao.jar"

EXTENSOES_PERMITIDAS = ('.doc', '.docx')

for pasta in [UPLOAD_FOLDER, PROCESSED_FOLDER, PRE_FICHAS_FOLDER, MODELS_FOLDER, BIN_FOLDER]:
    os.makedirs(pasta, exist_ok=True)

# --- FUNÇÕES AUXILIARES ---

def extrair_nome(caminho):
    try:
        doc = Document(caminho)
        texto_completo = " ".join(p.text for p in doc.paragraphs)
        for t in doc.tables:
            for r in t.rows:
                for c in r.cells:
                    texto_completo += " " + c.text

        match = re.search(
            r'(?:NOME|PACIENTE|NM)\s*[:\s_.]+\s*([A-ZÀ-Ú\s]{5,})(?:\n|\s{2,}|DATA|NASC|END)',
            texto_completo, re.IGNORECASE
        )
        if match:
            nome = " ".join(match.group(1).strip().split())
            return nome.title()
        return "NomeNaoIdentificado"
    except Exception:
        return "ErroLeitura"

def gerar_nome_seguro(nome_paciente, pasta_destino):
    data_hoje = datetime.now().strftime('%Y%m%d')
    nome_base = f"{data_hoje}_{nome_paciente.replace(' ', '_')}"

    nome_base = re.sub(r'[\\/*?:"<>|]', "", nome_base)

    caminho_final = os.path.join(pasta_destino, f"{nome_base}.docx")

    contador = 1
    while os.path.exists(caminho_final):
        caminho_final = os.path.join(pasta_destino, f"{nome_base}_{contador}.docx")
        contador += 1

    return os.path.basename(caminho_final)

def converter_doc_para_docx(caminho_doc):
    pasta = os.path.dirname(caminho_doc)
    try:
        subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "docx", caminho_doc, "--outdir", pasta],
            stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, timeout=30, check=True
        )
        novo_caminho = caminho_doc.rsplit('.', 1)[0] + '.docx'
        return novo_caminho if os.path.exists(novo_caminho) else None
    except (subprocess.CalledProcessError, FileNotFoundError, subprocess.TimeoutExpired):
        return None

def substituir_placeholders_py(doc, dados):
    for p in doc.paragraphs:
        for key, value in dados.items():
            if key in p.text:
                inline = p.runs
                for i in range(len(inline)):
                    if key in inline[i].text:
                        text = inline[i].text.replace(key, str(value))
                        inline[i].text = text

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for key, value in dados.items():
                        if key in p.text:
                            inline = p.runs
                            for i in range(len(inline)):
                                if key in inline[i].text:
                                    text = inline[i].text.replace(key, str(value))
                                    inline[i].text = text

# --- ROTAS PRINCIPAIS ---

@app.route('/', methods=['GET', 'POST'])
def index():
    enviados, erros = [], []

    if request.method == 'POST':
        arquivos = request.files.getlist('files')

        for f in arquivos:
            if not f.filename.lower().endswith(EXTENSOES_PERMITIDAS):
                erros.append(f"{f.filename}: Extensão não permitida")
                continue

            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                f.save(tmp.name)
                caminho_tmp = tmp.name

            caminho_docx = caminho_tmp
            if f.filename.lower().endswith('.doc'):
                caminho_docx = converter_doc_para_docx(caminho_tmp)
                if not caminho_docx:
                    erros.append(f"{f.filename}: Falha ao converter .doc para .docx")
                    os.remove(caminho_tmp)
                    continue

            nome_paciente = extrair_nome(caminho_docx)
            nome_final = gerar_nome_seguro(nome_paciente, UPLOAD_FOLDER)
            caminho_destino = os.path.join(UPLOAD_FOLDER, nome_final)

            shutil.move(caminho_docx, caminho_destino)
            enviados.append(f"{f.filename} -> {nome_final}")

            if f.filename.lower().endswith('.doc'):
                os.remove(caminho_tmp)

    return render_template('index.html', enviados=enviados, erros=erros)

@app.route('/historico')
def historico():
    pendentes, concluidos, em_processamento = [], [], []

    for f in sorted(os.listdir(PROCESSED_FOLDER), reverse=True):
        if f.endswith('.docx'):
            concluidos.append(f)

    for f in sorted(os.listdir(UPLOAD_FOLDER), reverse=True):
        if f.endswith('.docx'):
            caminho_completo = os.path.join(UPLOAD_FOLDER, f)
            stat = os.stat(caminho_completo)
            if stat.st_mtime > stat.st_ctime + 1:
                em_processamento.append(f)
            else:
                pendentes.append(f)

    return render_template('historico.html', pendentes=pendentes, concluidos=concluidos, em_processamento=em_processamento)

@app.route('/view/<path:filename>')
def view_file(filename):
    filename = unquote(filename)
    caminho = None

    caminho_pendente = os.path.join(UPLOAD_FOLDER, filename)
    caminho_concluido = os.path.join(PROCESSED_FOLDER, filename)

    if os.path.exists(caminho_pendente):
        caminho = caminho_pendente
    elif os.path.exists(caminho_concluido):
        caminho = caminho_concluido

    if not caminho:
        return "Arquivo não encontrado", 404

    try:
        doc = Document(caminho)
        texto_completo = []
        for p in doc.paragraphs:
            texto_completo.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                texto_linha = "\t".join(cell.text for cell in row.cells)
                texto_completo.append(texto_linha)

        return render_template('view.html', filename=filename, content="\n".join(texto_completo))
    except Exception as e:
        return f"Erro ao ler o arquivo: {e}", 500

# --- ROTAS PARA FICHAS PRÉ-PRONTAS (RASCUNHOS) ---

@app.route('/api/fichas/pre', methods=['GET', 'POST'])
def gerenciar_rascunhos():
    if request.method == 'POST':
        dados = request.get_json()
        if not dados or 'nome_rascunho' not in dados:
            return jsonify({"erro": "Nome do rascunho é obrigatório"}), 400

        nome_rascunho = re.sub(r'[\\/*?:"<>|]', "", dados['nome_rascunho'])
        data_agendamento = dados.get('data_agendamento', datetime.now().strftime('%Y-%m-%d'))

        nome_arquivo = f"{data_agendamento}_{nome_rascunho}.json"
        caminho_arquivo = os.path.join(PRE_FICHAS_FOLDER, nome_arquivo)

        contador = 1
        while os.path.exists(caminho_arquivo):
            caminho_arquivo = os.path.join(PRE_FICHAS_FOLDER, f"{data_agendamento}_{nome_rascunho}_{contador}.json")
            contador += 1

        try:
            with open(caminho_arquivo, 'w', encoding='utf-8') as f:
                json.dump(dados, f, ensure_ascii=False, indent=4)
            return jsonify({"sucesso": "Rascunho salvo com sucesso", "filename": os.path.basename(caminho_arquivo)}), 201
        except Exception as e:
            return jsonify({"erro": str(e)}), 500

    # GET
    try:
        data_filtro = request.args.get('data')
        arquivos = os.listdir(PRE_FICHAS_FOLDER)

        if data_filtro:
            arquivos_filtrados = [f for f in arquivos if f.startswith(data_filtro) and f.endswith('.json')]
        else:
            arquivos_filtrados = [f for f in arquivos if f.endswith('.json')]

        return jsonify(sorted(arquivos_filtrados, reverse=True))
    except Exception as e:
        return jsonify({"erro": str(e)}), 500

@app.route('/api/fichas/pre/<path:filename>', methods=['GET', 'DELETE'])
def gerenciar_rascunho_especifico(filename):
    filename = unquote(filename) # Decodifica o nome do arquivo
    caminho = os.path.join(PRE_FICHAS_FOLDER, filename)
    if not os.path.exists(caminho):
        return jsonify({"erro": "Rascunho não encontrado"}), 404

    if request.method == 'GET':
        return send_file(caminho, mimetype='application/json')

    if request.method == 'DELETE':
        try:
            os.remove(caminho)
            return jsonify({"sucesso": "Rascunho deletado"}), 200
        except Exception as e:
            return jsonify({"erro": str(e)}), 500

# --- ROTA PARA PACIENTE INTERNADO ---
@app.route('/api/laudos/internado', methods=['POST'])
def gerar_laudo_internado():
    dados_req = request.get_json()
    if not dados_req:
        return jsonify({"erro": "Requisição sem JSON"}), 400

    campos_necessarios = ["nome", "cns", "nasc", "procedencia", "tipo"]
    if not all(campo in dados_req for campo in campos_necessarios):
        return jsonify({"erro": "Dados incompletos no JSON"}), 400

    tipo_laudo = dados_req['tipo']
    nome_modelo = {
        'cateterismo': 'Laudo de cateterismo.docx',
        'angioplastia': 'Laudo de Angioplastia.docx'
    }.get(tipo_laudo)

    if not nome_modelo:
        return jsonify({"erro": f"Tipo de laudo inválido: {tipo_laudo}"}), 400

    caminho_modelo = os.path.join(MODELS_FOLDER, nome_modelo)
    if not os.path.exists(caminho_modelo):
        return jsonify({"erro": f"Arquivo de modelo não encontrado: {nome_modelo}"}), 500

    try:
        dados_laudo = {
            "{{NOME}}": dados_req.get("nome", "").upper(),
            "{{CNS}}": dados_req.get("cns", ""),
            "{{NASCIMENTO}}": dados_req.get("nasc", ""),
            "{{PROCEDENCIA}}": dados_req.get("procedencia", "").upper(),
            "{{DATA_HOJE}}": datetime.now().strftime('%d/%m/%Y'),
            "{{LOGRADOURO}}": "", "{{BAIRRO}}": "", "{{CIDADE}}": "", "{{UF}}": "",
            "{{RG}}": "", "{{CPF}}": "", "{{MAE}}": "", "{{PAI}}": "",
            "{{MATRICULA}}": "", "{{CHAVE}}": "", "{{TELEFONES}}": ""
        }

        doc = Document(caminho_modelo)
        substituir_placeholders_py(doc, dados_laudo)

        nome_final = gerar_nome_seguro(dados_req["nome"], UPLOAD_FOLDER)
        caminho_final = os.path.join(UPLOAD_FOLDER, nome_final)
        doc.save(caminho_final)

        return jsonify({"sucesso": "Laudo gerado com sucesso", "filename": nome_final}), 201

    except Exception as e:
        return jsonify({"erro": f"Falha ao gerar documento: {str(e)}"}), 500

# --- ROTAS DE API (LEGADAS) ---

@app.route('/api/pendentes')
def api_pendentes():
    lista = []
    for f in os.listdir(UPLOAD_FOLDER):
        if f.endswith('.docx'):
            lista.append({
                'arquivo': f,
                'nome': extrair_nome(os.path.join(UPLOAD_FOLDER, f))
            })
    return jsonify(lista)

@app.route('/api/historico')
def api_historico_legado():
    lista = []
    for f in os.listdir(PROCESSED_FOLDER):
        if f.endswith('.docx'):
            lista.append({'arquivo': f})
    return jsonify(lista)

@app.route('/api/excluir/<filename>', methods=['DELETE'])
def excluir_arquivo(filename):
    filename = unquote(filename)
    caminho_pendente = os.path.join(UPLOAD_FOLDER, filename)
    caminho_concluido = os.path.join(PROCESSED_FOLDER, filename)

    deletado = False
    if os.path.exists(caminho_pendente):
        os.remove(caminho_pendente)
        deletado = True
    if os.path.exists(caminho_concluido):
        os.remove(caminho_concluido)
        deletado = True

    if deletado: return "Arquivo excluído", 200
    return "Arquivo não encontrado", 404

@app.route('/api/concluir/<filename>', methods=['POST'])
def concluir(filename):
    filename = unquote(filename)
    origem = os.path.join(UPLOAD_FOLDER, filename)
    destino = os.path.join(PROCESSED_FOLDER, filename)
    if os.path.exists(origem):
        shutil.move(origem, destino)
        return "Concluído", 200
    return "Não encontrado", 404

@app.route('/api/baixar/<filename>')
def baixar(filename):
    filename = unquote(filename)
    for pasta in [UPLOAD_FOLDER, PROCESSED_FOLDER]:
        caminho = os.path.join(pasta, filename)
        if os.path.exists(caminho):
            return send_file(caminho, as_attachment=True)
    return "Arquivo não encontrado", 404

@app.route('/api/modelo/<tipo>')
def modelo(tipo):
    nome = {
        'cateterismo': 'Laudo de cateterismo.docx',
        'angioplastia': 'Laudo de Angioplastia.docx'
    }.get(tipo)
    if not nome: return "Tipo inválido", 400
    caminho = os.path.join(MODELS_FOLDER, nome)
    if os.path.exists(caminho): return send_file(caminho, as_attachment=True)
    return "Modelo não encontrado", 404

# --- ROTAS DE VERSÃO E ATUALIZAÇÃO ---

@app.route('/api/version')
def api_version_hemo():
    return jsonify({"version": APP_HEMO_VERSION, "filename": APP_HEMO_FILE})

@app.route('/api/download/app')
def download_hemo():
    caminho = os.path.join(BIN_FOLDER, APP_HEMO_FILE)
    if os.path.exists(caminho): return send_file(caminho, as_attachment=True)
    return "AppHemo.jar não encontrado na pasta bin", 404

@app.route('/api/recepcao/version')
def api_version_recepcao():
    return jsonify({"version": APP_RECEPCAO_VERSION, "filename": APP_RECEPCAO_FILE})

@app.route('/api/download/recepcao')
def download_recepcao():
    caminho = os.path.join(BIN_FOLDER, APP_RECEPCAO_FILE)
    if os.path.exists(caminho): return send_file(caminho, as_attachment=True)
    return "AppRecepcao.jar não encontrado na pasta bin", 404

# --- EXECUÇÃO ---
if __name__ == '__main__':
    app.run(host='0.0.0.0', port=2424, threaded=True)