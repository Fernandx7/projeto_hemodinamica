import os
import re
import shutil
import subprocess
import tempfile
import json
from datetime import datetime
from urllib.parse import unquote

from flask import Flask, request, jsonify, render_template, send_file
from docx import Document
import google.generativeai as genai

app = Flask(__name__)

# --- CONFIGURAÇÕES ---
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
UPLOAD_FOLDER = os.path.abspath(os.path.join(BASE_DIR, '..', 'uploads'))
WAITING_PROCEDURE_FOLDER = os.path.join(UPLOAD_FOLDER, 'aguardando_procedimento')
PROCESSED_FOLDER = os.path.join(UPLOAD_FOLDER, 'concluidos')
PRE_FICHAS_FOLDER = os.path.abspath(os.path.join(BASE_DIR, '..', 'pre_fichas'))
MODELS_FOLDER = os.path.abspath(os.path.join(BASE_DIR, 'teste_modelos'))
ATIVOS_FOLDER = os.path.abspath(os.path.join(BASE_DIR, '..', 'procedimentos_ativos'))
CHAVE_API_PATH = os.path.join(BASE_DIR, '..', 'laudo_falado', 'chaveapi')

EXTENSOES_PERMITIDAS = ('.doc', '.docx')

for pasta in [UPLOAD_FOLDER, WAITING_PROCEDURE_FOLDER, PROCESSED_FOLDER, PRE_FICHAS_FOLDER, MODELS_FOLDER, ATIVOS_FOLDER]:
    os.makedirs(pasta, exist_ok=True)

try:
    if os.path.exists(CHAVE_API_PATH):
        with open(CHAVE_API_PATH, "r") as f: genai.configure(api_key=f.read().strip())
except Exception: pass

# ------------------------------------------------------------------
# FUNÇÕES AUXILIARES
# ------------------------------------------------------------------

def extrair_dados_ia(texto_completo):
    try:
        model = genai.GenerativeModel("models/gemini-3.1-flash-lite-preview")
        prompt = f"Extraia Nome completo, CNS, Data de Nascimento e Procedência deste texto médico e retorne APENAS JSON puro: {texto_completo[:3500]}"
        response = model.generate_content(prompt)
        return json.loads(response.text.replace('```json', '').replace('```', '').strip())
    except Exception: return None

def extrair_dados_ficha(caminho):
    dados = {"nome": "NomeNaoIdentificado", "cns": "", "nasc": "", "procedencia": ""}
    if not caminho or not os.path.exists(caminho): return dados
    try:
        doc = Document(caminho)
        # Extrai texto de parágrafos e tabelas (inclui texto dentro de Content Controls)
        texto = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        for t in doc.tables:
            for r in t.rows:
                for c in r.cells: texto += "\n" + c.text
        
        d_ia = extrair_dados_ia(texto)
        if d_ia and d_ia.get("nome") and "NomeNaoIdentificado" not in d_ia["nome"]: return d_ia
        
        m = re.search(r'(?:NOME|PACIENTE|NM)\s*[:\s_.]+\s*([A-ZÀ-Úa-zà-ú\s]{5,100}?)(\n|\r|DATA|NASC|CNS|CPF|PROCED|ORIGEM|CONVÊNIO|\s{2,})', texto, re.I)
        if m: dados["nome"] = " ".join(m.group(1).strip().split('\n')[0].split()).title()
        
        m_proc = re.search(r'(?:PROCEDÊNCIA|ORIGEM|VINDO DE|CONVÊNIO)\s*[:\s_.]+\s*([A-ZÀ-Úa-zà-ú\s0-9]{3,})', texto, re.I)
        if m_proc: dados["procedencia"] = m_proc.group(1).strip().split('\n')[0].title()
        
        return dados
    except Exception: return dados

def gerar_laudo_final(nome, d_ext, tipo, materiais, num=None):
    sufixo = 'PTCA' if tipo == 'ANGIOPLASTIA' else 'CATETERISMO'
    modelo = 'Laudo de cateterismo.docx' if sufixo == 'CATETERISMO' else 'Laudo de Angioplastia.docx'
    path_m = os.path.join(MODELS_FOLDER, modelo)
    if not os.path.exists(path_m): return None
    try:
        d_sub = {
            "{{NOME}}": to_pascal_case(nome), "campo_nome": to_pascal_case(nome),
            "{{CNS}}": d_ext.get("cns", ""), "campo_cns": d_ext.get("cns", ""),
            "{{NASC}}": d_ext.get("nasc", ""), "{{NASCIMENTO}}": d_ext.get("nasc", ""), "campo_nasc": d_ext.get("nasc", ""),
            "{{PROCEDENCIA}}": to_pascal_case(d_ext.get("procedencia", "")), "campo_procedencia": to_pascal_case(d_ext.get("procedencia", "")),
            "{{DATA_HOJE}}": datetime.now().strftime('%d/%m/%Y'), "campo_data": datetime.now().strftime('%d/%m/%Y'),
            "{{NUM_EXAME}}": num if num else "{{NUM_EXAME}}", "campo_num_exame": num if num else "",
            "{{materiais}}": "\n".join([f"- {m}" for m in materiais]) if materiais else "Nenhum material.",
            "campo_materiais": "\n".join([f"- {m}" for m in materiais]) if materiais else ""
        }
        doc = Document(path_m)
        substituir_placeholders_py(doc, d_sub)
        nome_f = f"{datetime.now().strftime('%Y%m%d')}_{nome.replace(' ', '_')}_{sufixo}.docx"
        doc.save(os.path.join(UPLOAD_FOLDER, nome_f)); return nome_f
    except Exception: return None

def to_pascal_case(valor):
    if not valor or str(valor).upper() == "NOMENAOIDENTIFICADO": return ""
    partes = re.split(r"\s+", str(valor).strip())
    return " ".join(p[:1].upper() + p[1:].lower() for p in partes if p)

def substituir_placeholders_py(doc, dados):
    """Substitui mantendo estilo e penetrando em Controles de Conteúdo."""
    for p in doc.paragraphs:
        for k, v in dados.items():
            if k in p.text:
                new_t = p.text.replace(k, str(v))
                if p.runs:
                    for i in range(1, len(p.runs)): p.runs[i].text = ""
                    p.runs[0].text = new_t
                else: p.add_run(new_t)
                if "\n" in str(v): p.alignment = 0
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells: substituir_placeholders_py(cell, dados)

# ------------------------------------------------------------------
# ROTAS
# ------------------------------------------------------------------

@app.route('/', methods=['GET', 'POST'])
def index():
    enviados, erros = [], []
    if request.method == 'POST':
        for f in request.files.getlist('files'):
            try:
                with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                    f.save(tmp.name); d = extrair_dados_ficha(tmp.name)
                    nome_p = d['nome'] if d['nome'] != "NomeNaoIdentificado" else f.filename.split('.')[0]
                    n_f = f"{datetime.now().strftime('%Y%m%d')}_{nome_p.replace(' ', '_')}.docx"
                    shutil.move(tmp.name, os.path.join(WAITING_PROCEDURE_FOLDER, n_f))
                    enviados.append(f.filename)
            except Exception as e: erros.append(str(e))
    return render_template('index.html', enviados=enviados, erros=erros)

@app.route('/api/pendentes')
def api_pendentes():
    l = []
    if os.path.exists(UPLOAD_FOLDER):
        for f in os.listdir(UPLOAD_FOLDER):
            if f.endswith('.docx') and not os.path.isdir(os.path.join(UPLOAD_FOLDER, f)):
                d = extrair_dados_ficha(os.path.join(UPLOAD_FOLDER, f))
                l.append({"arquivo": f, "nome": d["nome"], "procedencia": d["procedencia"], "cns": d["cns"], "nasc": d["nasc"]})
    return jsonify(l)

@app.route('/api/sala/finalizar', methods=['POST'])
def finalizar_mesa():
    a = os.listdir(ATIVOS_FOLDER)
    if not a: return jsonify({"erro": "404"}), 404
    with open(os.path.join(ATIVOS_FOLDER, a[0]), 'r', encoding='utf-8') as f: info = json.load(f)
    nome_limpo = info['nome'].replace(' ', '_')
    ficha_path = None
    for f in os.listdir(WAITING_PROCEDURE_FOLDER):
        if nome_limpo in f: ficha_path = os.path.join(WAITING_PROCEDURE_FOLDER, f); break
    d_ext = extrair_dados_ficha(ficha_path) if ficha_path else {"nome": info['nome'], "procedencia": ""}
    num_cat = info.get('numero_exame')
    num_angio = None
    if num_cat and '.' in num_cat:
        try: p1, p2 = num_cat.split('.'); num_angio = f"{p1}.{int(p2)+1:03d}"
        except: pass
    gerar_laudo_final(info['nome'], d_ext, 'CATETERISMO', info.get('materiais_cateterismo', []), num_cat)
    if info.get('evoluiu_angioplastia'):
        gerar_laudo_final(info['nome'], d_ext, 'ANGIOPLASTIA', info.get('materiais_angioplastia', []), num_angio)
    if ficha_path: os.remove(ficha_path)
    os.remove(os.path.join(ATIVOS_FOLDER, a[0]))
    return jsonify({"status": "finalizado"})

# Mantenha as outras rotas... (omitidas aqui mas presentes no arquivo)
@app.route('/historico')
def historico():
    p = sorted(os.listdir(WAITING_PROCEDURE_FOLDER), reverse=True) if os.path.exists(WAITING_PROCEDURE_FOLDER) else []
    e = sorted([f for f in os.listdir(UPLOAD_FOLDER) if f.endswith('.docx') and not os.path.isdir(os.path.join(UPLOAD_FOLDER, f))], reverse=True)
    c = sorted(os.listdir(PROCESSED_FOLDER), reverse=True) if os.path.exists(PROCESSED_FOLDER) else []
    return render_template('historico.html', pendentes=p, em_processamento=e, concluidos=c)

@app.route('/tablet')
def tablet_index(): return render_template('tablet.html')

@app.route('/view/<path:filename>')
def view_file(filename):
    filename = unquote(filename)
    for p in [UPLOAD_FOLDER, PROCESSED_FOLDER, WAITING_PROCEDURE_FOLDER]:
        path = os.path.join(p, filename)
        if os.path.exists(path):
            doc = Document(path); return render_template('view.html', filename=filename, content="\n".join([p.text for p in doc.paragraphs]))
    return "404", 404

@app.route('/api/sala/espera')
def api_sala_espera():
    l = []
    if os.path.exists(WAITING_PROCEDURE_FOLDER):
        for f in os.listdir(WAITING_PROCEDURE_FOLDER):
            if f.endswith('.docx'):
                l.append({"arquivo": f, "nome": extrair_dados_ficha(os.path.join(WAITING_PROCEDURE_FOLDER, f))["nome"]})
    return jsonify(l)

@app.route('/api/sala/iniciar', methods=['POST'])
def iniciar_mesa():
    d = request.get_json(); n = d.get('nome')
    for f in os.listdir(ATIVOS_FOLDER): os.remove(os.path.join(ATIVOS_FOLDER, f))
    info = {"nome": n, "numero_exame": d.get('numero_exame'), "procedimento": "cateterismo", "inicio": datetime.now().strftime('%H:%M:%S'),
            "materiais_cateterismo": [], "materiais_angioplastia": [], "evoluiu_angioplastia": False}
    with open(os.path.join(ATIVOS_FOLDER, f"ativo.json"), 'w', encoding='utf-8') as f: json.dump(info, f, indent=4)
    return jsonify({"status": "iniciado"}), 201

@app.route('/api/sala/ativo')
def get_sala_ativo():
    a = os.listdir(ATIVOS_FOLDER)
    if not a: return jsonify({"status": "vazio"})
    with open(os.path.join(ATIVOS_FOLDER, a[0]), 'r', encoding='utf-8') as f: return jsonify(json.load(f))

@app.route('/api/sala/atualizar', methods=['POST'])
def atualizar_mesa():
    d = request.get_json(); a = os.listdir(ATIVOS_FOLDER)
    if not a: return jsonify({"erro": "404"}), 404
    with open(os.path.join(ATIVOS_FOLDER, a[0]), 'r', encoding='utf-8') as f: info = json.load(f)
    if 'material' in d:
        m = d['material'] + (" " + d.get('calibre', "")).strip()
        if info.get('evoluiu_angioplastia'): info['materiais_angioplastia'].append(m)
        else: info['materiais_cateterismo'].append(m)
    if 'evoluiu_angioplastia' in d: info['evoluiu_angioplastia'] = d['evoluiu_angioplastia']
    with open(os.path.join(ATIVOS_FOLDER, a[0]), 'w', encoding='utf-8') as f: json.dump(info, f, indent=4)
    return jsonify({"status": "OK"})

@app.route('/api/materiais', methods=['GET'])
def get_materiais():
    if not os.path.exists(MATERIAIS_FILE): return jsonify({"catalogo": {}})
    with open(MATERIAIS_FILE, 'r', encoding='utf-8') as f: data = json.load(f)
    if request.args.get('procedimento') == 'cateterismo': data['catalogo'] = {"acesso_e_diagnostico": data['catalogo'].get('acesso_e_diagnostico', [])}
    return jsonify(data)

@app.route('/api/baixar/<path:filename>')
def baixar(filename):
    for p in [UPLOAD_FOLDER, PROCESSED_FOLDER, WAITING_PROCEDURE_FOLDER]:
        path = os.path.join(p, unquote(filename))
        if os.path.exists(path): return send_file(path, as_attachment=True)
    return "404", 404

@app.route('/api/concluir/<path:filename>', methods=['POST'])
def concluir(filename):
    origem = os.path.join(UPLOAD_FOLDER, unquote(filename))
    if os.path.exists(origem): shutil.move(origem, os.path.join(PROCESSED_FOLDER, unquote(filename))); return "OK", 200
    return "404", 404

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=2400, threaded=True)
