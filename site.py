import os
import re
import shutil
import subprocess
import tempfile
import json
from datetime import datetime
from urllib.parse import unquote

from flask import Flask, request, jsonify, render_template, send_file, redirect, url_for
from docx import Document

app = Flask(__name__)

# --- CONFIGURAÇÕES ---
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
UPLOAD_FOLDER = os.path.join(BASE_DIR, 'uploads')
PROCESSED_FOLDER = os.path.join(UPLOAD_FOLDER, 'concluidos')
PRE_FICHAS_FOLDER = os.path.join(BASE_DIR, 'pre_fichas')
MODELS_FOLDER = os.path.join(BASE_DIR, 'modelos')
BIN_FOLDER = os.path.join(BASE_DIR, 'bin')
MATERIAIS_FILE = os.path.join(BASE_DIR, 'materiais.json')
ATIVOS_FOLDER = os.path.join(BASE_DIR, 'procedimentos_ativos')

EXTENSOES_PERMITIDAS = ('.doc', '.docx')

for pasta in [UPLOAD_FOLDER, PROCESSED_FOLDER, PRE_FICHAS_FOLDER, MODELS_FOLDER, BIN_FOLDER, ATIVOS_FOLDER]:
    os.makedirs(pasta, exist_ok=True)

# ------------------------------------------------------------------
# FUNÇÕES AUXILIARES
# ------------------------------------------------------------------

def extrair_dados_ficha(caminho):
    """
    Tenta extrair Nome, CNS, Nascimento e Procedência da ficha carregada.
    """
    dados = {
        "nome": "NomeNaoIdentificado",
        "cns": "",
        "nasc": "",
        "procedencia": ""
    }
    try:
        doc = Document(caminho)
        texto_completo = " ".join(p.text for p in doc.paragraphs)
        for t in doc.tables:
            for r in t.rows:
                for c in r.cells:
                    texto_completo += " " + c.text

        # Nome
        match_nome = re.search(
            r'(?:NOME|PACIENTE|NM)\s*[:\s_.]+\s*([A-ZÀ-Úa-zà-ú\s]{5,})(?:\n|\s{2,}|DATA|NASC|END|CNS)',
            texto_completo, re.IGNORECASE
        )
        if match_nome:
            dados["nome"] = " ".join(match_nome.group(1).strip().split()).title()

        # CNS
        match_cns = re.search(r'(?:CNS|CARTÃO SUS)\s*[:\s_.]+\s*(\d{5,})', texto_completo, re.IGNORECASE)
        if match_cns:
            dados["cns"] = match_cns.group(1).strip()

        # Nascimento
        match_nasc = re.search(r'(?:NASC|NASCIMENTO|DATA DE NASC)\s*[:\s_.]+\s*(\d{2}/\d{2}/\d{4})', texto_completo, re.IGNORECASE)
        if match_nasc:
            dados["nasc"] = match_nasc.group(1).strip()

        # Procedência
        match_proc = re.search(r'(?:PROCEDÊNCIA|ORIGEM|VINDO DE)\s*[:\s_.]+\s*([A-ZÀ-Úa-zà-ú\s0-9]{3,})(?:\n|\s{2,}|DATA|LEITO|CNS)', texto_completo, re.IGNORECASE)
        if match_proc:
            dados["procedencia"] = match_proc.group(1).strip().title()

        return dados
    except Exception:
        return dados


def gerar_laudos_automaticos(nome_paciente, dados_extracao):
    """
    Gera as versões pré-prontas de Cateterismo e Angioplastia para o paciente.
    """
    tipos = {
        'CATETERISMO': 'Laudo de cateterismo.docx',
        'ANGIOPLASTIA': 'Laudo de Angioplastia.docx'
    }

    gerados = []
    for tipo_label, modelo_nome in tipos.items():
        caminho_modelo = os.path.join(MODELS_FOLDER, modelo_nome)
        if not os.path.exists(caminho_modelo):
            continue

        try:
            dados_laudo = {
                "{{NOME}}": to_pascal_case(nome_paciente),
                "{{CNS}}": dados_extracao.get("cns", ""),
                "{{NASC}}": dados_extracao.get("nasc", ""),
                "{{NASCIMENTO}}": dados_extracao.get("nasc", ""),
                "{{PROCEDENCIA}}": to_pascal_case(dados_extracao.get("procedencia", "")),
                "{{DATA_HOJE}}": datetime.now().strftime('%d/%m/%Y'),
                "{{NUM_EXAME}}": "{{NUM_EXAME}}", # Mantém o placeholder para o Java
                "{{LOGRADOURO}}": "", "{{BAIRRO}}": "", "{{CIDADE}}": "", "{{UF}}": "",
                "{{RG}}": "", "{{CPF}}": "", "{{MAE}}": "", "{{PAI}}": "",
                "{{MATRICULA}}": "", "{{CHAVE}}": "", "{{TELEFONES}}": ""
            }

            doc = Document(caminho_modelo)
            substituir_placeholders_py(doc, dados_laudo)

            # Nome do arquivo: DATA_NOME_TIPO.docx
            data_str = datetime.now().strftime('%Y%m%d')
            nome_limpo = re.sub(r'[\\/*?:"<>|]', "", nome_paciente.replace(' ', '_'))
            nome_arquivo = f"{data_str}_{nome_limpo}_{tipo_label}.docx"

            # Evita sobrescrever se já existir
            caminho_final = os.path.join(UPLOAD_FOLDER, nome_arquivo)
            contador = 1
            while os.path.exists(caminho_final):
                caminho_final = os.path.join(UPLOAD_FOLDER, f"{data_str}_{nome_limpo}_{tipo_label}_{contador}.docx")
                contador += 1

            doc.save(caminho_final)
            gerados.append(os.path.basename(caminho_final))
        except Exception as e:
            print(f"Erro ao gerar {tipo_label}: {e}")

    return gerados


def extrair_nome(caminho):
    dados = extrair_dados_ficha(caminho)
    return dados["nome"]


def gerar_nome_seguro(nome_paciente, pasta_destino):
    data = datetime.now().strftime('%Y%m%d')
    nome_base = f"{data}_{nome_paciente.replace(' ', '_')}"

    nome_base = re.sub(r'[\\/*?:"<>|]', "", nome_base)
    caminho = os.path.join(pasta_destino, f"{nome_base}.docx")
    contador = 1

    while os.path.exists(caminho):
        caminho = os.path.join(pasta_destino, f"{nome_base}_{contador}.docx")
        contador += 1

    return os.path.basename(caminho)


def converter_doc_para_docx(caminho_doc):
    pasta = os.path.dirname(caminho_doc)
    try:
        subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "docx", caminho_doc, "--outdir", pasta],
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
            timeout=30,
            check=True
        )
        novo = caminho_doc.rsplit('.', 1)[0] + '.docx'
        return novo if os.path.exists(novo) else None
    except Exception:
        return None


def to_pascal_case(valor):
    if not valor:
        return ""
    partes = re.split(r"\s+", valor.strip())
    return " ".join(p[:1].upper() + p[1:].lower() for p in partes if p)


def substituir_placeholders_py(doc, dados):
    """
    Percorre paragrafos e tabelas do documento e substitui os placeholders.
    """
    for p in doc.paragraphs:
        for key, value in dados.items():
            if key in p.text:
                for i in range(len(p.runs)):
                    if key in p.runs[i].text:
                        p.runs[i].text = p.runs[i].text.replace(key, value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for key, value in dados.items():
                        if key in p.text:
                            for i in range(len(p.runs)):
                                if key in p.runs[i].text:
                                    p.runs[i].text = p.runs[i].text.replace(key, value)

# ------------------------------------------------------------------
# ROTAS WEB
# ------------------------------------------------------------------

@app.route('/', methods=['GET', 'POST'])
def index():
    enviados, erros = [], []

    if request.method == 'POST':
        arquivos = request.files.getlist('files')

        for f in arquivos:
            if not f.filename.lower().endswith(EXTENSOES_PERMITIDAS):
                erros.append(f"{f.filename}: Extensão não permitida")
                continue

            caminho_original = None
            caminho_docx = None

            try:
                sufixo = '.doc' if f.filename.lower().endswith('.doc') else '.docx'
                with tempfile.NamedTemporaryFile(delete=False, suffix=sufixo) as tmp:
                    f.save(tmp.name)
                    caminho_original = tmp.name

                caminho_docx = caminho_original

                if f.filename.lower().endswith('.doc'):
                    caminho_docx = converter_doc_para_docx(caminho_original)
                    if not caminho_docx:
                        erros.append(f"{f.filename}: Falha ao converter .doc para .docx")
                        continue

                dados_ficha = extrair_dados_ficha(caminho_docx)
                nome_paciente = dados_ficha["nome"]
                nome_final = gerar_nome_seguro(nome_paciente, UPLOAD_FOLDER)
                caminho_destino = os.path.join(UPLOAD_FOLDER, nome_final)

                shutil.move(caminho_docx, caminho_destino)
                
                # Gera laudos automáticos (COMENTADO POR SOLICITAÇÃO)
                # gerados = gerar_laudos_automaticos(nome_paciente, dados_ficha)
                
                msg = f"{f.filename} → {nome_final}"
                # if gerados:
                #     msg += f" (Laudos gerados: {', '.join(gerados)})"
                enviados.append(msg)

            except Exception as e:
                erros.append(f"{f.filename}: {str(e)}")

            finally:
                for c in [caminho_original, caminho_docx]:
                    try:
                        if c and os.path.exists(c) and not c.startswith(UPLOAD_FOLDER):
                            os.remove(c)
                    except FileNotFoundError:
                        pass

    return render_template('index.html', enviados=enviados, erros=erros)


@app.route('/historico')
def historico():
    pendentes, concluidos, em_processamento = [], [], []

    for f in sorted(os.listdir(PROCESSED_FOLDER), reverse=True):
        if f.endswith('.docx'):
            concluidos.append(f)

    for f in sorted(os.listdir(UPLOAD_FOLDER), reverse=True):
        if f.endswith('.docx'):
            caminho = os.path.join(UPLOAD_FOLDER, f)
            stat = os.stat(caminho)
            if stat.st_mtime > stat.st_ctime + 1:
                em_processamento.append(f)
            else:
                pendentes.append(f)

    return render_template(
        'historico.html',
        pendentes=pendentes,
        concluidos=concluidos,
        em_processamento=em_processamento
    )


@app.route('/view/<path:filename>')
def view_file(filename):
    filename = unquote(filename)
    for pasta in [UPLOAD_FOLDER, PROCESSED_FOLDER]:
        caminho = os.path.join(pasta, filename)
        if os.path.exists(caminho):
            doc = Document(caminho)
            texto = []

            for p in doc.paragraphs:
                texto.append(p.text)

            for t in doc.tables:
                for r in t.rows:
                    texto.append("\t".join(c.text for c in r.cells))

            return render_template('view.html', filename=filename, content="\n".join(texto))

    return "Arquivo não encontrado", 404

# ------------------------------------------------------------------
# ROTAS API
# ------------------------------------------------------------------

# --- FICHAS PRÉ-PRONTAS (RASCUNHOS) ---

@app.route('/api/fichas/pre', methods=['POST'])
def salvar_rascunho():
    dados = request.get_json()
    if not dados or 'nome_rascunho' not in dados:
        return jsonify({"erro": "Nome do rascunho é obrigatório"}), 400

    nome_rascunho = re.sub(r'[\\/*?:"<>|]', "", dados['nome_rascunho'])
    data_hoje = datetime.now().strftime('%Y-%m-%d')
    nome_arquivo = f"{data_hoje}_{nome_rascunho}.json"

    caminho_arquivo = os.path.join(PRE_FICHAS_FOLDER, nome_arquivo)

    # Evita sobrescrever
    contador = 1
    while os.path.exists(caminho_arquivo):
        caminho_arquivo = os.path.join(PRE_FICHAS_FOLDER, f"{data_hoje}_{nome_rascunho}_{contador}.json")
        contador += 1

    try:
        with open(caminho_arquivo, 'w', encoding='utf-8') as f:
            json.dump(dados, f, ensure_ascii=False, indent=4)
        return jsonify({"sucesso": "Rascunho salvo com sucesso", "filename": os.path.basename(caminho_arquivo)}), 201
    except Exception as e:
        return jsonify({"erro": str(e)}), 500

@app.route('/api/fichas/pre', methods=['GET'])
def listar_rascunhos():
    try:
        # Lista e ordena por data (mais recentes primeiro)
        arquivos = sorted(os.listdir(PRE_FICHAS_FOLDER), reverse=True)
        return jsonify([f for f in arquivos if f.endswith('.json')])
    except Exception as e:
        return jsonify({"erro": str(e)}), 500

@app.route('/api/fichas/pre/<path:filename>', methods=['GET'])
def carregar_rascunho(filename):
    caminho = os.path.join(PRE_FICHAS_FOLDER, filename)
    if not os.path.exists(caminho):
        return jsonify({"erro": "Rascunho não encontrado"}), 404
    return send_file(caminho, mimetype='application/json')

@app.route('/api/fichas/pre/<path:filename>', methods=['DELETE'])
def deletar_rascunho(filename):
    caminho = os.path.join(PRE_FICHAS_FOLDER, filename)
    if not os.path.exists(caminho):
        return jsonify({"erro": "Rascunho não encontrado"}), 404
    try:
        os.remove(caminho)
        return jsonify({"sucesso": "Rascunho deletado"}), 200
    except Exception as e:
        return jsonify({"erro": str(e)}), 500


@app.route('/api/laudos/internado', methods=['POST'])
def gerar_laudo_internado():
    """
    Recebe dados de um paciente internado e gera AMBOS os laudos (CAT e ANGIO)
    na pasta 'uploads', deixando-os prontos para o Java preencher o número.
    """
    dados_req = request.get_json()
    if not dados_req:
        return jsonify({"erro": "Requisição sem JSON"}), 400

    campos_necessarios = ["nome", "cns", "nasc", "procedencia"]
    if not all(campo in dados_req for campo in campos_necessarios):
        return jsonify({"erro": "Dados incompletos no JSON"}), 400

    try:
        nome_paciente = dados_req.get("nome", "")
        # Reutiliza a função de geração automática
        gerados = gerar_laudos_automaticos(nome_paciente, dados_req)

        if not gerados:
            return jsonify({"erro": "Nenhum laudo foi gerado. Verifique os modelos."}), 500

        return jsonify({
            "sucesso": "Laudos gerados com sucesso",
            "filenames": gerados
        }), 201

    except Exception as e:
        return jsonify({"erro": f"Falha ao gerar documentos: {str(e)}"}), 500


@app.route('/api/pendentes')
def api_pendentes():
    lista = []
    for f in os.listdir(UPLOAD_FOLDER):
        if f.endswith('.docx'):
            lista.append({
                "arquivo": f,
                "nome": extrair_nome(os.path.join(UPLOAD_FOLDER, f))
            })
    return jsonify(lista)


@app.route('/api/historico')
def api_historico():
    lista = []
    for f in os.listdir(PROCESSED_FOLDER):
        if f.endswith('.docx'):
            lista.append({"arquivo": f})
    return jsonify(lista)


@app.route('/api/excluir/<filename>', methods=['DELETE'])
def excluir_arquivo(filename):
    filename = unquote(filename)
    deletado = False

    for pasta in [UPLOAD_FOLDER, PROCESSED_FOLDER]:
        caminho = os.path.join(pasta, filename)
        if os.path.exists(caminho):
            os.remove(caminho)
            deletado = True

    return ("Arquivo excluído", 200) if deletado else ("Arquivo não encontrado", 404)


@app.route('/api/concluir/<filename>', methods=['POST'])
def concluir(filename):
    filename = unquote(filename)
    origem = os.path.join(UPLOAD_FOLDER, filename)
    destino = os.path.join(PROCESSED_FOLDER, filename)

    if os.path.exists(origem):
        shutil.move(origem, destino)
        return "Concluído", 200

    return "Não encontrado", 404


@app.route('/api/baixar/<filename>')
def baixar(filename):
    filename = unquote(filename)

    for pasta in [UPLOAD_FOLDER, PROCESSED_FOLDER]:
        caminho = os.path.join(pasta, filename)
        if os.path.exists(caminho):
            return send_file(caminho, as_attachment=True)

    return "Arquivo não encontrado", 404


@app.route('/excluir/<filename>', methods=['POST'])
def excluir_arquivo_post(filename):
    filename = unquote(filename)
    
    sucesso = False
    for pasta in [UPLOAD_FOLDER, PROCESSED_FOLDER]:
        caminho = os.path.join(pasta, filename)
        if os.path.exists(caminho):
            try:
                os.remove(caminho)
                sucesso = True
            except Exception as e:
                print(f"Erro ao excluir {caminho}: {e}")
    
    if sucesso:
        return redirect(url_for('historico'))
    
    return "Arquivo não encontrado ou erro ao excluir", 404


@app.route('/excluir_multiplos', methods=['POST'])
def excluir_multiplos():
    filenames = request.form.getlist('filenames')
    
    if not filenames:
        return redirect(url_for('historico'))
        
    for filename in filenames:
        filename = unquote(filename)
        for pasta in [UPLOAD_FOLDER, PROCESSED_FOLDER]:
            caminho = os.path.join(pasta, filename)
            if os.path.exists(caminho):
                try:
                    os.remove(caminho)
                except Exception as e:
                    print(f"Erro ao excluir {caminho}: {e}")
    
    return redirect(url_for('historico'))


@app.route('/api/modelo/<tipo>')
def modelo(tipo):
    nome = {
        'cateterismo': 'Laudo de cateterismo.docx',
        'angioplastia': 'Laudo de Angioplastia.docx'
    }.get(tipo)

    if not nome:
        return "Tipo inválido", 400

    caminho = os.path.join(MODELS_FOLDER, nome)
    return send_file(caminho, as_attachment=True) if os.path.exists(caminho) else ("Modelo não encontrado", 404)


# --- ROTAS DA SALA (TABLET) ---

@app.route('/tablet')
def tablet_index():
    return render_template('tablet.html')

@app.route('/api/materiais', methods=['GET'])
def get_materiais():
    if not os.path.exists(MATERIAIS_FILE):
        return jsonify({"configuracoes": {}, "catalogo": {}})
    with open(MATERIAIS_FILE, 'r', encoding='utf-8') as f:
        materiais_data = json.load(f)
    
    procedimento = request.args.get('procedimento')
    if procedimento == 'cateterismo':
        # Filtra apenas o que é pertinente ao cateterismo
        materiais_data['catalogo'] = {
            "acesso_e_diagnostico": materiais_data['catalogo'].get('acesso_e_diagnostico', [])
        }
    return jsonify(materiais_data)

@app.route('/api/materiais', methods=['POST'])
def save_material():
    """Adiciona ou atualiza um material no catálogo visualmente."""
    novo_item = request.get_json()
    tipo_cat = novo_item.get('categoria') 
    
    if not os.path.exists(MATERIAIS_FILE):
        materiais_data = {
            "configuracoes": {"kit_padrao_cateterismo": []},
            "catalogo": {
                "acesso_e_diagnostico": [],
                "intervencao": {"stents": [], "baloes": [], "guias": []}
            }
        }
    else:
        with open(MATERIAIS_FILE, 'r', encoding='utf-8') as f:
            materiais_data = json.load(f)

    catalogo = materiais_data.get('catalogo', {})
    intervencao = catalogo.get('intervencao', {})

    if tipo_cat == 'acesso_e_diagnostico':
        catalogo['acesso_e_diagnostico'].append(novo_item)
    elif tipo_cat in ['stents', 'baloes', 'guias']:
        if tipo_cat in ['stents', 'baloes']:
            marca = novo_item.get('marca')
            found = False
            for item in intervencao.get(tipo_cat, []):
                if item['marca'] == marca and item.get('tipo') == novo_item.get('tipo'):
                    item['medidas'] = list(set(item.get('medidas', []) + novo_item.get('medidas', [])))
                    found = True
                    break
            if not found:
                intervencao.setdefault(tipo_cat, []).append(novo_item)
        else:
            intervencao.setdefault(tipo_cat, []).append(novo_item)
    else:
        return jsonify({"erro": "Categoria inválida"}), 400

    with open(MATERIAIS_FILE, 'w', encoding='utf-8') as f:
        json.dump(materiais_data, f, indent=4, ensure_ascii=False)
    return jsonify({"status": "sucesso"}), 201

@app.route('/api/sala/ativo', methods=['GET'])
def get_paciente_em_mesa():
    """Retorna o paciente que está atualmente em procedimento, se houver."""
    arquivos = os.listdir(ATIVOS_FOLDER)
    if not arquivos:
        return jsonify({"status": "vazio"})
    
    with open(os.path.join(ATIVOS_FOLDER, arquivos[0]), 'r', encoding='utf-8') as f:
        return jsonify(json.load(f))

@app.route('/api/sala/iniciar', methods=['POST'])
def iniciar_mesa():
    """Inicia um procedimento para um paciente selecionado."""
    dados = request.get_json()
    nome_paciente = dados.get('nome')
    procedimento = dados.get('procedimento', 'cateterismo')
    
    # Limpa procedimentos ativos anteriores
    for f in os.listdir(ATIVOS_FOLDER):
        os.remove(os.path.join(ATIVOS_FOLDER, f))

    info = {
        "nome": nome_paciente,
        "procedimento": procedimento,
        "inicio": datetime.now().strftime('%H:%M:%S'),
        "materiais_cateterismo": [],
        "materiais_angioplastia": [],
        "evoluiu_angioplastia": False
    }

    # Se for Cateterismo, carrega o Kit Padrão na lista de cateterismo
    if procedimento == 'cateterismo':
        if os.path.exists(MATERIAIS_FILE):
            with open(MATERIAIS_FILE, 'r', encoding='utf-8') as f:
                materiais_data = json.load(f)
                kit = materiais_data.get('configuracoes', {}).get('kit_padrao_cateterismo', [])
                for item in kit:
                    nome_item = item['nome']
                    if 'calibre_padrao' in item:
                        nome_item += f" {item['calibre_padrao']}"
                    info['materiais_cateterismo'].append(nome_item)
    
    filename = f"ativo_{re.sub(r'[^a-zA-Z0-9]', '_', nome_paciente)}.json"
    with open(os.path.join(ATIVOS_FOLDER, filename), 'w', encoding='utf-8') as f:
        json.dump(info, f, indent=4)
        
    return jsonify({"status": "iniciado"}), 201

@app.route('/api/sala/atualizar', methods=['POST'])
def atualizar_mesa():
    """Adiciona material ou muda status (CAT -> ANGIO)."""
    dados = request.get_json()
    arquivos = os.listdir(ATIVOS_FOLDER)
    if not arquivos:
        return jsonify({"erro": "Nenhum paciente em mesa"}), 404

    caminho = os.path.join(ATIVOS_FOLDER, arquivos[0])
    with open(caminho, 'r', encoding='utf-8') as f:
        info = json.load(f)

    if 'material' in dados:
        material_nome = dados['material']
        calibre = dados.get('calibre')
        if calibre:
            material_nome = f"{material_nome} {calibre}"
        
        # Define em qual lista o material deve entrar
        if info.get('evoluiu_angioplastia'):
            info['materiais_angioplastia'].append(material_nome)
        else:
            info['materiais_cateterismo'].append(material_nome)
        
    if 'evoluiu_angioplastia' in dados:
        info['evoluiu_angioplastia'] = dados['evoluiu_angioplastia']

    with open(caminho, 'w', encoding='utf-8') as f:
        json.dump(info, f, indent=4)
        
    return jsonify({"status": "atualizado"})

@app.route('/api/sala/finalizar', methods=['POST'])
def finalizar_mesa():
    """Finaliza o procedimento e limpa a sala."""
    arquivos = os.listdir(ATIVOS_FOLDER)
    if not arquivos:
        return jsonify({"erro": "Nenhum paciente em mesa"}), 404

    # Aqui poderíamos salvar um log histórico se necessário
    for f in os.listdir(ATIVOS_FOLDER):
        os.remove(os.path.join(ATIVOS_FOLDER, f))
        
    return jsonify({"status": "finalizado"})

# ------------------------------------------------------------------

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=2424, threaded=True)