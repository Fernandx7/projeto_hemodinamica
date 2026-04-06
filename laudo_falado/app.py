import os
import time
import shutil
import re
from fastapi import FastAPI, Request, HTTPException, File, UploadFile, Form
from fastapi.responses import HTMLResponse, FileResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import google.generativeai as genai
from pydub import AudioSegment

app = FastAPI()

# Configurações
CHAVE_API_PATH = "chaveapi"
RESULTS_DIR = "results"
TEMP_DIR = "temp_uploads"
MODELS_DIR = "../atualização mega/teste_modelos" # Aponta para a pasta de teste

if not os.path.exists(RESULTS_DIR): os.makedirs(RESULTS_DIR)
if not os.path.exists(TEMP_DIR): os.makedirs(TEMP_DIR)

try:
    if os.path.exists(CHAVE_API_PATH):
        with open(CHAVE_API_PATH, "r") as f: genai.configure(api_key=f.read().strip())
except Exception as e: print(f"Erro Chave IA: {e}")

app.mount("/static", StaticFiles(directory="static"), name="static")
templates = Jinja2Templates(directory="templates")

def formatar_arial_11(paragraph, texto):
    """Insere texto em um parágrafo com formatação Arial 11."""
    paragraph.text = "" # Limpa rascunho
    run = paragraph.add_run(texto)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

def preencher_modelo_ia_v2(template_name, texto_ia, caminho_saida):
    """
    Abre o modelo _ditado, preenche a parte clínica em Arial 11.
    """
    try:
        caminho_modelo = os.path.join(MODELS_DIR, template_name)
        if not os.path.exists(caminho_modelo):
            # Tenta na pasta local se não achar na modelos
            caminho_modelo = template_name 
            
        doc = Document(caminho_modelo)
        texto_limpo = texto_ia.replace("**", "").strip()
        
        # Como o modelo agora só tem títulos, vamos inserir o texto 
        # nos locais apropriados ou ao final das seções.
        # Estratégia: Encontrar o parágrafo que contém o título e inserir o texto abaixo.
        
        # Para simplificar e garantir Arial 11 em tudo que a IA gerar:
        substituiu = False
        for p in doc.paragraphs:
            if any(m in p.text for m in ["{{laudo}}", "{{CONTEUDO}}", "CONCLUSÃO:", "Conclusão:"]):
                formatar_arial_11(p, p.text.replace("{{laudo}}", "").replace("{{CONTEUDO}}", "") + "\n" + texto_limpo)
                substituiu = True
                break
        
        if not substituido:
            # Se não achou marcador, adiciona novo parágrafo Arial 11 no fim
            p = doc.add_paragraph()
            formatar_arial_11(p, texto_limpo)

        doc.save(caminho_saida)
        return True
    except Exception as e:
        print(f"Erro IA Filler V2: {e}")
        return False

@app.post("/process")
async def process_report(
    audio_file: UploadFile = File(...),
    template_name: str = Form("Laudo de cateterismo_ditado.docx")
):
    temp_p = os.path.join(TEMP_DIR, audio_file.filename)
    mp3_p = os.path.join(TEMP_DIR, f"conv_{int(time.time())}.mp3")
    try:
        with open(temp_p, "wb") as b: shutil.copyfileobj(audio_file.file, b)
        try:
            AudioSegment.converter = "/usr/bin/ffmpeg"
            AudioSegment.ffprobe = "/usr/bin/ffprobe"
            audio = AudioSegment.from_file(temp_p)
            audio.export(mp3_p, format="mp3")
            proc_p = mp3_p
        except Exception: proc_p = temp_p

        model = genai.GenerativeModel("models/gemini-3.1-flash-lite-preview")
        sample = genai.upload_file(path=proc_p)
        while sample.state.name == "PROCESSING":
            time.sleep(2)
            sample = genai.get_file(sample.name)
            
        prompt = """
        Você é um cardiologista intervencionista. 
        Transcreva o áudio para o corpo do laudo médico.
        REGRAS:
        1. Use linguagem formal.
        2. NÃO use Markdown (sem asteriscos).
        3. NÃO repita o cabeçalho do hospital.
        4. Organize em seções se o médico ditar (Ex: Descendente Anterior, Circunflexa, etc).
        """
        
        response = model.generate_content([prompt, sample])
        out_name = f"Gerado_{int(time.time())}.docx"
        out_path = os.path.join(RESULTS_DIR, out_name)
        
        preencher_modelo_ia_v2(template_name, response.text, out_path)
        
        genai.delete_file(sample.name)
        return {"status": "success", "download_url": f"/download/{out_name}"}
    except Exception as e:
        return {"status": "error", "message": str(e)}

@app.get("/download/{filename}")
async def download_file(filename: str):
    return FileResponse(os.path.join(RESULTS_DIR, filename), filename=filename)

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
